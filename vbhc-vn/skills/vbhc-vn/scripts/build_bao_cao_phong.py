#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_bao_cao_phong.py — Dựng BÁO CÁO ĐỊNH KỲ CỦA PHÒNG (tháng / quý / 6 tháng / 9 tháng / năm)
trên mẫu thật `examples/sct/bao-cao-thang-phong-qlcn.docx` (Chế độ B) và tự chuẩn hóa thể thức.

Đúc kết từ ngày 06/9/2026 (Báo cáo 9 tháng, Báo cáo tháng 9 của Phòng QLCN):
  1. Giữ nguyên header 2 cột + khối ký "TRƯỞNG PHÒNG Nguyễn Hữu Long" của mẫu; chỉ thay tiêu đề và thân.
  2. Header chuẩn NĐ 30: cơ quan chủ quản 12pt thường; tên Phòng 13pt ĐẬM + Line; Quốc hiệu 12pt ĐẬM
     (để nằm gọn 1 dòng trong cột 5800 twips); Tiêu ngữ 13pt ĐẬM + Line; dòng ngày 13pt nghiêng, ngày để trống.
  3. Thân 14pt căn đều, thụt đầu dòng 1,25 cm đồng nhất, giãn dòng ĐƠN (KHÔNG dùng "Exactly" —
     Word tính khác LibreOffice, gây trống nửa trang và lệch số trang), cách đoạn sau 3pt, widowControl.
  4. Đề mục I/II/III và 1./2. đậm; a)/b) nghiêng không đậm; khối ký cantSplit.
  5. Sau khi dựng: chạy qa_all.py; chữ lẻ rơi dòng (WIDOW) sửa bằng --widow "<cụm cuối đoạn>".

CÁCH DÙNG
    python3 build_bao_cao_phong.py noidung.txt out.docx \
        --title1 "Tình hình thực hiện công tác tháng 9 năm 2026," \
        --title2 "Kế hoạch công tác tháng 10 năm 2026" \
        [--month 10] [--year 2026] [--base examples/sct/bao-cao-thang-phong-qlcn.docx] [--widow "cụm 1" "cụm 2"]

File nội dung: mỗi dòng một đoạn, thẻ đầu dòng:
    [H]  đề mục đậm (I. / 1. / II. ...)      [I] đề mục nghiêng (a) / b))      [P] hoặc không thẻ: đoạn thường
Dòng bắt đầu '##' là chú thích, bỏ qua. KHÔNG dùng **, *, markdown trong nội dung.
"""
import argparse, os, re, sys, zipfile
from copy import deepcopy
from docx import Document

HERE = os.path.dirname(os.path.abspath(__file__))
DEFAULT_BASE = os.path.join(HERE, '..', 'examples', 'sct', 'bao-cao-thang-phong-qlcn.docx')


def read_content(path):
    rows = []
    for line in open(path, encoding='utf-8'):
        line = line.rstrip('\n')
        if not line.strip() or line.startswith('##'):
            continue
        kind = 'b'
        for tag, k in (('[H]', 'h'), ('[I]', 'i'), ('[P]', 'b')):
            if line.startswith(tag):
                kind = k
                line = line[len(tag):].strip()
                break
        assert '**' not in line and '\n' not in line, f'không dùng markdown/ngắt dòng cứng: {line[:60]}'
        rows.append((line, kind))
    return rows


def set_text(p_el, text, bold=False, italic=False):
    from docx.text.paragraph import Paragraph
    p = Paragraph(p_el, None)
    runs = p.runs
    assert runs, 'paragraph mẫu phải có run'
    runs[0].text = text
    runs[0].bold = True if bold else None
    runs[0].italic = True if italic else None
    for r in runs[1:]:
        r._element.getparent().remove(r._element)


def build_body(base, out, title1, title2, rows, month=None, year=None):
    d = Document(base)
    ps = d.paragraphs
    if month or year:
        for p in d.tables[0].rows[0].cells[1].paragraphs:
            if 'Lào Cai, ngày' in p.text:
                m = re.search(r'tháng\s*(\d+)\s*năm\s*(\d{4})', p.text)
                mo = month or m.group(1); yr = year or m.group(2)
                p.runs[0].text = f'Lào Cai, ngày      tháng {mo} năm {yr}'
                for r in p.runs[1:]:
                    r._element.getparent().remove(r._element)
    # Tìm vị trí: p[1]='BÁO CÁO', p[2], p[3] tiêu đề; thân từ đề mục "I." đến trước 2 đoạn trống cuối
    i_first = next(i for i, p in enumerate(ps) if p.text.strip().startswith('I.'))
    i_last = len(ps) - 1
    while i_last > i_first and not ps[i_last].text.strip():
        i_last -= 1
    set_text(ps[2]._p, title1, bold=True)
    set_text(ps[3]._p, title2, bold=True)
    model_head = next(deepcopy(p._p) for p in ps[i_first:] if p.runs and p.runs[0].bold and not p.text.startswith('I.'))
    model_body = next(deepcopy(p._p) for p in ps[i_first:] if p.runs and not p.runs[0].bold and len(p.text) > 80)
    old = [p._p for p in ps[i_first:i_last + 1]]
    anchor = ps[i_first - 1]._p
    new = []
    for text, kind in rows:
        el = deepcopy(model_head if kind == 'h' else model_body)
        set_text(el, text, bold=(kind == 'h'), italic=(kind == 'i'))
        new.append(el)
    for el in reversed(new):
        anchor.addnext(el)
    for el in old:
        el.getparent().remove(el)
    d.save(out)


def rw(path, fn):
    with zipfile.ZipFile(path) as z:
        x = z.read('word/document.xml').decode('utf-8')
    x = fn(x)
    tmp = path + '.tmp'
    with zipfile.ZipFile(path) as zin, zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:
        for it in zin.infolist():
            data = zin.read(it.filename)
            if it.filename == 'word/document.xml':
                data = x.encode('utf-8')
            zout.writestr(it, data)
    os.replace(tmp, path)


def _run_of(x, text):
    m = re.search(r'<w:r\b[^>]*>(?:(?!</w:r>).)*?' + re.escape(text) + r'(?:(?!</w:r>).)*?</w:r>', x, re.S)
    assert m, f'không thấy run chứa: {text}'
    return m


def _set_run(x, text, size=None, bold=None):
    m = _run_of(x, text)
    run = m.group(0)
    if size:
        run = re.sub(r'<w:sz w:val="\d+"/>', f'<w:sz w:val="{size}"/>', run)
        run = re.sub(r'<w:szCs w:val="\d+"/>', f'<w:szCs w:val="{size}"/>', run)
        if '<w:sz ' not in run:
            run = run.replace('<w:rPr>', f'<w:rPr><w:sz w:val="{size}"/><w:szCs w:val="{size}"/>', 1)
    if bold and '<w:b/>' not in run:
        run = run.replace('<w:rPr>', '<w:rPr><w:b/>', 1) if '<w:rPr>' in run else run.replace('<w:r>', '<w:r><w:rPr><w:b/></w:rPr>', 1)
    return x[:m.start()] + run + x[m.end():]


def vml_line(width_pt, n):
    return ('<w:r><w:rPr><w:noProof/></w:rPr><w:pict>'
            '<v:line xmlns:v="urn:schemas-microsoft-com:vml" xmlns:o="urn:schemas-microsoft-com:office:office" '
            f'id="HeaderRule{n}" o:spid="_x0000_s10{25 + n}" '
            f'style="position:absolute;z-index:2516582{39 + n};mso-position-horizontal:center;'
            'mso-position-horizontal-relative:column;mso-position-vertical-relative:line" '
            f'from="0,23pt" to="{width_pt}pt,23pt" o:gfxdata="" strokecolor="black" strokeweight=".75pt"/></w:pict></w:r>')


def fix_format(x, widows):
    # Quốc hiệu chuẩn
    x = x.replace('CỘNG HOÀ', 'CỘNG HÒA').replace('Độc lập - Tự do - Hạnh phúc', 'Độc lập – Tự do – Hạnh phúc')
    # Header: cỡ chữ + đậm + Line
    i0, i1 = x.find('<w:tbl>'), x.find('</w:tbl>')
    hdr = x[i0:i1]
    hdr = hdr.replace('<w:sz w:val="28"/>', '<w:sz w:val="26"/>').replace('<w:szCs w:val="28"/>', '<w:szCs w:val="26"/>')
    hdr = hdr.replace('<w:tblW w:w="9214" w:type="dxa"/>', '<w:tblW w:w="9400" w:type="dxa"/>')
    hdr = hdr.replace('<w:gridCol w:w="3686"/><w:gridCol w:w="5528"/>', '<w:gridCol w:w="3600"/><w:gridCol w:w="5800"/>')
    hdr = hdr.replace('<w:tcW w:w="3686" w:type="dxa"/>', '<w:tcW w:w="3600" w:type="dxa"/>').replace('<w:tcW w:w="5528" w:type="dxa"/>', '<w:tcW w:w="5800" w:type="dxa"/>')
    x = x[:i0] + hdr + x[i1:]
    if '<w:pict' not in x[i0:x.find('</w:tbl>')]:
        for n, (anchor, w) in enumerate([('PHÒNG QL CÔNG NGHIỆP', 110), ('Độc lập – Tự do – Hạnh phúc', 130)], 1):
            idx = x.find(anchor); end = x.find('</w:p>', idx)
            x = x[:end] + vml_line(w, n) + x[end:]
    x = _set_run(x, 'SỞ CÔNG THƯƠNG LÀO CAI', size=24)
    x = _set_run(x, 'PHÒNG QL CÔNG NGHIỆP', bold=True)
    x = _set_run(x, 'CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM', size=24, bold=True)
    x = _set_run(x, 'Độc lập – Tự do – Hạnh phúc', bold=True)
    # Thân: giãn dòng đơn, cách đoạn 3pt, widowControl
    b0 = x.find('</w:tbl>') + 8; b1 = x.rfind('<w:tbl>')
    body = re.sub(r'<w:spacing w:line="\d+" w:lineRule="exact"/>',
                  '<w:widowControl/><w:spacing w:before="0" w:after="60" w:line="240" w:lineRule="auto"/>', x[b0:b1])
    # bỏ 1 đoạn trống thừa trước khối ký
    body = re.sub(r'(<w:p\b[^>]*>(?:(?!<w:t[ >]).)*?</w:p>)\s*$', '', body, count=1, flags=re.S)
    x = x[:b0] + body + x[b1:]
    # Khối ký không gãy trang
    seg = x[x.rfind('<w:tbl>'):]
    if '<w:cantSplit/>' not in seg:
        seg = re.sub(r'<w:trPr>', '<w:trPr><w:cantSplit/>', seg, count=1)
        if '<w:cantSplit/>' not in seg:
            seg = re.sub(r'(<w:tr\b[^>]*>)', r'\1<w:trPr><w:cantSplit/></w:trPr>', seg, count=1)
        x = x[:x.rfind('<w:tbl>')] + seg
    # Chữ lẻ rơi dòng: co khoảng cách chữ -4 trong run chứa cụm cuối đoạn
    for key in widows:
        m = _run_of(x, key)
        run = m.group(0)
        if '<w:spacing' not in run:
            run = run.replace('<w:rPr>', '<w:rPr><w:spacing w:val="-4"/>', 1)
        x = x[:m.start()] + run + x[m.end():]
    return x


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('content'); ap.add_argument('out')
    ap.add_argument('--title1', required=True); ap.add_argument('--title2', required=True)
    ap.add_argument('--base', default=DEFAULT_BASE)
    ap.add_argument('--widow', nargs='*', default=[])
    ap.add_argument('--month'); ap.add_argument('--year')
    a = ap.parse_args()
    rows = read_content(a.content)
    build_body(a.base, a.out, a.title1, a.title2, rows, a.month, a.year)
    rw(a.out, lambda x: fix_format(x, a.widow))
    print(f'Đã dựng {a.out} ({len(rows)} đoạn). Tiếp theo: python3 qa_all.py "{a.out}" --require "{a.title1[:20]}"')


if __name__ == '__main__':
    main()
