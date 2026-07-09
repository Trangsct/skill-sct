#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
qa_all.py — QA MỘT PHÁT cho văn bản hành chính (.docx), thay chuỗi lệnh rời rạc
của Bước 4 cũ (validate → render soi ảnh → qa_pdf_check → check_document),
mỗi thứ một lượt tool và render PDF trùng lặp.

Một lệnh duy nhất làm trọn:
  1. Kiểm XML không cần render: Line header, 13pt dòng Số/Ngày (từ qa_pdf_check),
     <w:br/> trong header (Quy tắc bất biến 10), đoạn body căn giữa / thiếu
     firstLine 1cm (WARN — bài học Chế độ B).
  2. check_document.py: VBQPPL hết hiệu lực, từ suy đoán, số văn bản đáng ngờ.
  3. Render PDF ĐÚNG MỘT LẦN (profile soffice ấm — nhanh hơn từ lần chạy 2)
     rồi dùng lại cho: widow word + khối ký gãy trang (qa_pdf_check),
     xuất ảnh từng trang, và GHÉP TẤT CẢ TRANG THÀNH 1 ẢNH qa_sheet.png
     — Claude chỉ cần MỘT lượt `view` thay vì xem từng trang.
  4. Báo cáo gọn PASS/FAIL từng mục + exit code (0 = PASS, 1 = có FAIL).

Cách dùng:
  python3 qa_all.py <file.docx> [--min-lines N] [--no-image] [--dpi 100]

Ảnh ra tại: /home/claude/work/qa/qa_sheet.png (ảnh ghép) + qa-N.jpg (từng trang,
chỉ soi riêng khi ảnh ghép phát hiện nghi vấn).
Yêu cầu: soffice, pdftotext, pdftoppm, python-docx, Pillow (thiếu Pillow thì
bỏ qua ảnh ghép, vẫn xuất ảnh từng trang).
"""

import re
import subprocess
import sys
import time
from pathlib import Path

SCRIPT_DIR = Path(__file__).resolve().parent
sys.path.insert(0, str(SCRIPT_DIR))

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("Cần python-docx: pip install python-docx --break-system-packages")
    sys.exit(2)

# Tái dùng các hàm kiểm đã kiểm chứng của qa_pdf_check — KHÔNG viết lại
from qa_pdf_check import (  # noqa: E402
    check_line_shapes, check_so_ngay_13pt, check_widow, check_sig_split,
    pdf_pages_text,
)

QA_DIR = Path("/home/claude/work/qa")
LO_PROFILE = "file:///home/claude/.lo_profile"  # profile ấm: lần render sau nhanh hơn
CM1_EMU_TWIPS = 567  # 1cm ≈ 567 twips (w:firstLine tính bằng twips)


def render_pdf_once(docx_path: Path, outdir: Path) -> Path:
    outdir.mkdir(parents=True, exist_ok=True)
    subprocess.run(
        ["soffice", "--headless", f"-env:UserInstallation={LO_PROFILE}",
         "--convert-to", "pdf", "--outdir", str(outdir), str(docx_path)],
        check=True, capture_output=True, timeout=180,
    )
    pdf = outdir / (docx_path.stem + ".pdf")
    if not pdf.exists():
        raise RuntimeError("Render PDF thất bại")
    return pdf


def check_header_br(docx_path: Path):
    """Quy tắc bất biến 10: không <w:br/> trong ô header (table 0)."""
    issues = []
    doc = Document(str(docx_path))
    if not doc.tables:
        return issues
    header = doc.tables[0]
    n_br = sum(len(list(p._p.iter(qn("w:br"))))
               for row in header.rows for cell in row.cells
               for p in cell.paragraphs)
    if n_br:
        issues.append(f"header (table 0) chứa {n_br} thẻ <w:br/> — "
                      "V/v, tên cơ quan, ngày tháng phải là chuỗi liền, để Word tự wrap")
    return issues


def check_body_format(docx_path: Path):
    """WARN (bài học Chế độ B): đoạn body dài mà căn giữa, hoặc thiếu firstLine ≥ 1cm."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
    warns = []
    doc = Document(str(docx_path))
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if len(text) < 100:  # chỉ soi đoạn giống body; tiêu đề/ký/Quốc hiệu được miễn
            continue
        if p.alignment == AL.CENTER:
            warns.append(f"P{i} dài {len(text)} ký tự nhưng CĂN GIỮA "
                         f"(nghi clone nhầm từ đoạn tiêu đề): '{text[:50]}…'")
        fl = p.paragraph_format.first_line_indent
        if fl is not None and int(fl.twips if hasattr(fl, 'twips') else fl) < 0:
            warns.append(f"P{i} dùng thụt treo (hanging indent) — cấm theo thể thức: "
                         f"'{text[:50]}…'")
        elif fl is not None and 0 < int(fl.twips if hasattr(fl, 'twips') else fl) < CM1_EMU_TWIPS - 30:
            warns.append(f"P{i} firstLine < 1cm: '{text[:50]}…'")
    return warns


def run_check_document(docx_path: Path):
    """Chạy check_document.py (VBQPPL hết hiệu lực, từ suy đoán…), gom output."""
    cd = SCRIPT_DIR / "check_document.py"
    if not cd.exists():
        return None, ["không tìm thấy check_document.py — bỏ qua"]
    r = subprocess.run([sys.executable, str(cd), str(docx_path)],
                       capture_output=True, text=True, timeout=120)
    out = (r.stdout + r.stderr).strip()
    return r.returncode, [ln for ln in out.splitlines() if ln.strip()]


def make_contact_sheet(pdf: Path, dpi: int):
    """Xuất ảnh từng trang + ghép mọi trang thành 1 ảnh qa_sheet.png (lưới 3 cột)."""
    QA_DIR.mkdir(parents=True, exist_ok=True)
    for old in QA_DIR.glob("qa-*.jpg"):
        old.unlink()
    subprocess.run(["pdftoppm", "-jpeg", "-r", str(dpi), str(pdf),
                    str(QA_DIR / "qa")], check=True, capture_output=True, timeout=120)
    pages = sorted(QA_DIR.glob("qa-*.jpg"))
    sheet = None
    try:
        from PIL import Image, ImageDraw
        cols = min(3, len(pages))
        thumb_w = 720 if cols >= 3 else (900 if cols == 2 else 1100)
        imgs = []
        for pth in pages:
            im = Image.open(pth)
            h = int(im.height * thumb_w / im.width)
            imgs.append(im.resize((thumb_w, h)))
        rows = (len(imgs) + cols - 1) // cols
        pad, label_h = 10, 26
        cell_h = max(i.height for i in imgs) + label_h
        W = cols * (thumb_w + pad) + pad
        H = rows * (cell_h + pad) + pad
        canvas = Image.new("RGB", (W, H), "#666")
        draw = ImageDraw.Draw(canvas)
        for k, im in enumerate(imgs):
            x = pad + (k % cols) * (thumb_w + pad)
            y = pad + (k // cols) * (cell_h + pad)
            draw.text((x + 4, y + 4), f"— Trang {k + 1}/{len(imgs)} —", fill="#fff")
            canvas.paste(im, (x, y + label_h))
        sheet = QA_DIR / "qa_sheet.png"
        canvas.save(sheet)
    except ImportError:
        pass  # thiếu Pillow: vẫn còn ảnh từng trang
    return pages, sheet


def _docx_all_text_nfc(docx_path):
    """Toàn bộ text của .docx (paragraph + mọi ô bảng), chuẩn hóa NFC."""
    import unicodedata
    from docx import Document
    d = Document(str(docx_path))
    parts = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            for c in row.cells:
                parts.append(c.text)
    return unicodedata.normalize("NFC", "\n".join(parts))


def check_content_lists(docx_path, forbid, require):
    """
    Đối chiếu NỘI DUNG theo danh sách (bài học 09/7/2026 — vụ V/v nổ mìn:
    thể thức PASS nhưng trích yếu vẫn là vụ cũ vì replace trượt im lặng).
    - forbid: chuỗi vụ cũ KHÔNG được còn (tên DN, trích yếu, CN(tên) cũ...)
    - require: chuỗi vụ mới BẮT BUỘC có (số CV đến, tên người ký, CN(tên)...)
    So khớp sau chuẩn hóa NFC hai phía.
    """
    import unicodedata
    msgs = []
    text = _docx_all_text_nfc(docx_path)
    for w in forbid:
        if unicodedata.normalize("NFC", w) in text:
            msgs.append(f"còn dấu vết vụ cũ / chuỗi cấm: {w!r}")
    for w in require:
        if unicodedata.normalize("NFC", w) not in text:
            msgs.append(f"thiếu nội dung bắt buộc: {w!r}")
    return msgs


def _collect_list_arg(args, flag):
    """Gom giá trị sau --forbid/--require tới flag kế tiếp; loại khỏi args."""
    vals = []
    while flag in args:
        i = args.index(flag)
        args.pop(i)
        while i < len(args) and not args[i].startswith("--"):
            vals.append(args.pop(i))
    return vals


def main():
    t0 = time.time()
    args = sys.argv[1:]
    forbid_list = _collect_list_arg(args, "--forbid")
    require_list = _collect_list_arg(args, "--require")
    if not args:
        print(__doc__)
        sys.exit(2)
    docx_path = Path(args[0])
    if not docx_path.exists():
        print(f"LỖI: không tìm thấy file: {docx_path}")
        sys.exit(2)
    if docx_path.suffix.lower() != ".docx":
        print(f"LỖI: nhận .docx, không nhận '{docx_path.suffix}'. "
              "PDF thì dùng extract_metadata.py / pdftotext -layout.")
        sys.exit(2)
    min_lines = int(args[args.index("--min-lines") + 1]) if "--min-lines" in args else 2
    dpi = int(args[args.index("--dpi") + 1]) if "--dpi" in args else 100
    do_image = "--no-image" not in args

    fails, warns = [], []

    # ── 1. Kiểm XML (không cần render) ────────────────────────────────
    for msg in check_line_shapes(docx_path, min_lines):
        fails.append(("LINES", msg))
    for msg in check_so_ngay_13pt(docx_path):
        fails.append(("SZ13", msg))
    for msg in check_header_br(docx_path):
        fails.append(("HDR-BR", msg))
    for msg in check_body_format(docx_path):
        warns.append(("BODY", msg))
    if forbid_list or require_list:
        for msg in check_content_lists(docx_path, forbid_list, require_list):
            fails.append(("CONTENT", msg))

    # ── 2. check_document.py (nội dung: hiệu lực VBQPPL, từ suy đoán…) ─
    cd_code, cd_lines = run_check_document(docx_path)
    if cd_code not in (None, 0):
        fails.append(("NOIDUNG", "check_document.py phát hiện cảnh báo — xem mục dưới"))

    # ── 3. Render PDF một lần, dùng lại cho mọi kiểm + ảnh ────────────
    pages_txt, pages_img, sheet, pdf = None, [], None, None
    try:
        pdf = render_pdf_once(docx_path, QA_DIR)
        pages_txt = pdf_pages_text(pdf)
    except Exception as e:
        warns.append(("RENDER", f"không render được PDF: {e}"))
    if pages_txt:
        for msg in check_widow(pages_txt):
            fails.append(("WIDOW", msg))
        for msg in check_sig_split(pages_txt):
            fails.append(("SIGSPLIT", msg))
    if pdf and do_image:
        try:
            pages_img, sheet = make_contact_sheet(pdf, dpi)
        except Exception as e:
            warns.append(("IMAGE", f"không xuất được ảnh QA: {e}"))

    # ── 4. Báo cáo gọn ─────────────────────────────────────────────────
    print("=" * 62)
    print(f"QA MỘT PHÁT — {docx_path.name}")
    print("=" * 62)
    for tag, msg in warns:
        print(f"[WARN {tag}] {msg}")
    for tag, msg in fails:
        print(f"[FAIL {tag}] {msg}")
    if cd_lines:
        print("\n--- check_document.py ---")
        for ln in cd_lines[:40]:
            print("  " + ln)
        if len(cd_lines) > 40:
            print(f"  … ({len(cd_lines) - 40} dòng nữa — chạy riêng để xem đủ)")
    if pages_img:
        n_pages = len(pages_img)
    elif pages_txt:
        n_pages = len([p for p in pages_txt if p.strip()])  # bỏ trang rỗng do formfeed cuối
    else:
        n_pages = "?"
    print("\n--- Ảnh QA ---")
    if sheet:
        print(f"ẢNH GHÉP ({n_pages} trang trong 1 ảnh): {sheet}")
        print("→ Claude: `view` MỘT ảnh này là đủ soi tổng thể; chỉ mở qa-N.jpg "
              "khi ảnh ghép phát hiện nghi vấn ở trang N.")
    elif pages_img:
        print("Ảnh từng trang: " + ", ".join(str(p) for p in pages_img))
    verdict = "PASS" if not fails else f"FAIL ({len(fails)} lỗi)"
    print(f"\nKẾT QUẢ: {verdict}  |  thời gian: {time.time() - t0:.1f}s  |  "
          f"trang: {n_pages}")
    if fails:
        print("Sửa theo báo cáo TEXT ở trên trước, render lại chỉ khi đã sửa xong.")
    sys.exit(0 if not fails else 1)


if __name__ == "__main__":
    main()
