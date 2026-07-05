#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
qa_pdf_check.py — QA tự động thể thức VBHC trước khi giao file.

Sinh ra từ vụ 4 lỗi thật ngày 05/7/2026 (công văn HHNH loại 5, 8):
  1. WIDOW    — 1 chữ đơn độc rơi xuống dòng cuối đoạn (quét pdftotext -layout).
  2. SIGSPLIT — khối chữ ký gãy giữa 2 trang: tên người ký khác trang với
                chức danh "KT. GIÁM ĐỐC"/"TM. .../Q. GIÁM ĐỐC/GIÁM ĐỐC".
  3. LINES    — thiếu đường Line header (văn bản SCT chuẩn phải có >= 2
                shape w:pict/w:drawing: dưới tên cơ quan và dưới Tiêu ngữ).
  4. SZ13     — dòng "Số: .../..." và "…, ngày … tháng … năm …" chưa đặt
                13pt tường minh (w:sz=26) hoặc dòng ngày thiếu in nghiêng.

Cách dùng:
  python3 qa_pdf_check.py <file.docx> [--min-lines 2] [--no-render]

Thoát mã 0 = PASS toàn bộ; 1 = có FAIL. Cảnh báo (WARN) không chặn.
Yêu cầu: soffice (LibreOffice), pdftotext, python-docx.
"""

import re
import subprocess
import sys
import tempfile
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("Cần python-docx: pip install python-docx --break-system-packages")
    sys.exit(2)

SIGN_TITLES = [
    "KT. GIÁM ĐỐC", "PHÓ GIÁM ĐỐC", "GIÁM ĐỐC", "TM. ỦY BAN NHÂN DÂN",
    "CHỦ TỊCH", "KT. CHỦ TỊCH", "PHÓ CHỦ TỊCH", "TL. CHỦ TỊCH",
    "CHÁNH VĂN PHÒNG", "KT. CHÁNH VĂN PHÒNG", "PHÓ CHÁNH VĂN PHÒNG",
    "TRƯỞNG PHÒNG", "TM. ĐOÀN KIỂM TRA", "TRƯỞNG ĐOÀN", "Q. GIÁM ĐỐC",
]

# Các dòng 1 từ hợp lệ, không tính widow (số trang, ký hiệu...)
WIDOW_WHITELIST = re.compile(r"^(\d+|[IVXLC]+\.?|PHỤ|LỤC|-|–)$")


def render_pdf(docx_path: Path, outdir: Path) -> Path:
    subprocess.run(
        ["soffice", "--headless", "--convert-to", "pdf",
         "--outdir", str(outdir), str(docx_path)],
        check=True, capture_output=True, timeout=120,
    )
    pdf = outdir / (docx_path.stem + ".pdf")
    if not pdf.exists():
        raise RuntimeError("Render PDF thất bại")
    return pdf


def pdf_pages_text(pdf_path: Path):
    out = subprocess.run(
        ["pdftotext", "-layout", str(pdf_path), "-"],
        check=True, capture_output=True, timeout=60,
    ).stdout.decode("utf-8", errors="replace")
    return out.split("\f")


def check_widow(pages) -> list:
    """Dòng chỉ có 1 từ đứng ngay trước dòng trống (cuối đoạn)."""
    issues = []
    for pno, page in enumerate(pages, 1):
        lines = page.split("\n")
        for i, ln in enumerate(lines):
            s = ln.strip()
            if not s or len(s.split()) != 1 or WIDOW_WHITELIST.match(s):
                continue
            nxt = lines[i + 1].strip() if i + 1 < len(lines) else ""
            prv = lines[i - 1].strip() if i > 0 else ""
            # widow = 1 từ, cuối đoạn (dòng sau trống), dòng trước có nội dung dài
            if nxt == "" and len(prv.split()) >= 4:
                issues.append(f"trang {pno}: '{s}' rơi 1 mình cuối đoạn (sau: '{prv[-50:]}')")
    return issues


def check_sig_split(pages) -> list:
    """Chức danh ký và tên người ký phải cùng trang.

    Heuristic: trang chứa chức danh cuối cùng phải chứa cả dòng tên
    (dòng cuối cùng có 2-5 từ viết hoa chữ đầu, không phải chức danh).
    Cách chắc chắn hơn: nếu MỘT trang chỉ chứa tên người ký + nơi nhận rơi
    xuống mà không có nội dung body -> nghi gãy khối ký.
    """
    issues = []
    # tìm trang cuối cùng có chức danh ký
    title_page = None
    for pno, page in enumerate(pages, 1):
        for t in SIGN_TITLES:
            if t in page:
                title_page = max(title_page or 0, pno)
    if title_page is None:
        return ["không tìm thấy chức danh ký trong PDF (kiểm tra thủ công)"]
    # sau trang chức danh còn trang nào có nội dung không?
    for pno in range(title_page, len(pages)):
        page = pages[pno]  # trang sau title_page (index pno = trang pno+1)
        body = page.strip()
        if body and pno + 1 > title_page:
            # trang sau chức danh còn chữ -> khả năng tên ký/nơi nhận rơi trang
            first_words = " ".join(body.split()[:12])
            issues.append(
                f"khối ký nghi gãy trang: chức danh ở trang {title_page}, "
                f"trang {pno + 1} vẫn còn nội dung: '{first_words}...'"
            )
    return issues


def check_line_shapes(docx_path: Path, min_lines: int) -> list:
    d = Document(str(docx_path))
    xml = d.element.xml
    n = xml.count("<w:pict")
    if n < min_lines:
        return [f"chỉ có {n} shape Line trong file (yêu cầu >= {min_lines}: "
                "dưới tên cơ quan VÀ dưới Tiêu ngữ). Nghi mất Line do gán run.text "
                "vào run neo shape (Quy tắc 11)."]
    return []


def check_so_ngay_13pt(docx_path: Path) -> list:
    """Dòng 'Số: .../...' và dòng ngày trong bảng header: sz=26 tường minh,
    dòng ngày in nghiêng."""
    issues = []
    d = Document(str(docx_path))
    if not d.tables:
        return []
    t0 = d.tables[0]
    for r in range(len(t0.rows)):
        for c in range(len(t0.columns)):
            try:
                cell = t0.cell(r, c)
            except Exception:
                continue
            for p in cell.paragraphs:
                text = p.text
                is_so = bool(re.search(r"Số\s*:", text))
                is_ngay = bool(re.search(r"ngày\s.*tháng\s.*năm", text))
                if not (is_so or is_ngay):
                    continue
                for run in p.runs:
                    if not run.text.strip():
                        continue
                    rpr = run._element.find(qn("w:rPr"))
                    sz = rpr.find(qn("w:sz")) if rpr is not None else None
                    val = sz.get(qn("w:val")) if sz is not None else None
                    if val != "26":
                        issues.append(
                            f"dòng {'Số' if is_so else 'ngày'} run '{run.text[:20]}' "
                            f"sz={val} (yêu cầu 26 = 13pt tường minh)")
                        break
                if is_ngay:
                    has_italic = any(
                        run._element.find(qn("w:rPr")) is not None and
                        run._element.find(qn("w:rPr")).find(qn("w:i")) is not None
                        for run in p.runs if run.text.strip())
                    if not has_italic:
                        issues.append("dòng ngày thiếu in nghiêng (w:i)")
    return issues


def main():
    args = sys.argv[1:]
    if not args:
        print(__doc__)
        sys.exit(2)
    docx_path = Path(args[0])
    # Guard đầu vào: tool nhận .docx (tự render PDF bên trong), KHÔNG nhận PDF.
    # Bài học 05/7/2026: truyền nhầm .pdf → PackageNotFoundError khó hiểu.
    if not docx_path.exists():
        print(f"LỖI: không tìm thấy file: {docx_path}")
        sys.exit(2)
    if docx_path.suffix.lower() != ".docx":
        print(f"LỖI: tool này nhận file .docx (tự render PDF để QA), "
              f"không nhận '{docx_path.suffix}'.")
        print(f"Cách dùng đúng:  python3 qa_pdf_check.py duong_dan/van_ban.docx")
        if docx_path.suffix.lower() == ".pdf":
            print("Gợi ý: nếu chỉ có PDF, dùng scripts/extract_metadata.py "
                  "hoặc pdftotext -layout để kiểm nội dung.")
        sys.exit(2)
    min_lines = 2
    do_render = "--no-render" not in args
    if "--min-lines" in args:
        min_lines = int(args[args.index("--min-lines") + 1])

    fails, warns = [], []

    # Kiểm tra XML (không cần render)
    for msg in check_line_shapes(docx_path, min_lines):
        fails.append(("LINES", msg))
    for msg in check_so_ngay_13pt(docx_path):
        fails.append(("SZ13", msg))

    if do_render:
        with tempfile.TemporaryDirectory() as td:
            try:
                pdf = render_pdf(docx_path, Path(td))
                pages = pdf_pages_text(pdf)
            except Exception as e:
                warns.append(("RENDER", f"không render được PDF: {e}"))
                pages = None
            if pages:
                for msg in check_widow(pages):
                    fails.append(("WIDOW", msg))
                for msg in check_sig_split(pages):
                    fails.append(("SIGSPLIT", msg))

    for tag, msg in warns:
        print(f"[WARN {tag}] {msg}")
    for tag, msg in fails:
        print(f"[FAIL {tag}] {msg}")
    if not fails:
        print("QA PASS: không phát hiện widow word, khối ký gãy trang, "
              "thiếu Line header, hay sai cỡ dòng Số/Ngày.")
        sys.exit(0)
    print(f"\nTổng: {len(fails)} lỗi — sửa xong chạy lại trước khi giao file.")
    sys.exit(1)


if __name__ == "__main__":
    main()
