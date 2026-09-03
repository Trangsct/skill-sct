# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RED = RGBColor(0xC0, 0x00, 0x00)
IND = Cm(1.0)

doc = Document()
s = doc.sections[0]
s.page_width, s.page_height = Cm(21.0), Cm(29.7)
s.left_margin, s.right_margin = Cm(3.0), Cm(2.0)
s.top_margin, s.bottom_margin = Cm(2.0), Cm(2.0)

st = doc.styles['Normal']
st.font.name = 'Times New Roman'
st.font.size = Pt(14)
st.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
st.paragraph_format.space_after = Pt(0)
st.paragraph_format.space_before = Pt(0)
st.paragraph_format.line_spacing = 1.2
st.paragraph_format.alignment = AL.JUSTIFY


def add_run(p, text, bold=False, italic=False, size=14, red=False):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    if red:
        r.font.color.rgb = RED
    return r


def para(text='', bold=False, italic=False, align=AL.JUSTIFY, indent=IND,
         size=14, red=False, space_before=0, space_after=0):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = align
    p.paragraph_format.first_line_indent = indent
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        add_run(p, text, bold, italic, size, red)
    return p


def mixed(parts, indent=IND, align=AL.JUSTIFY, space_before=0, space_after=0):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = align
    p.paragraph_format.first_line_indent = indent
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    for txt, red in parts:
        add_run(p, txt, red=red)
    return p


def set_fixed(tbl, widths):
    lay = OxmlElement('w:tblLayout')
    lay.set(qn('w:type'), 'fixed')
    tbl._tbl.tblPr.append(lay)
    grid = tbl._tbl.find(qn('w:tblGrid'))
    if grid is not None:
        tbl._tbl.remove(grid)
    grid = OxmlElement('w:tblGrid')
    for w in widths:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(int(w.twips)))
        grid.append(gc)
    tbl._tbl.insert(1, grid)


def set_borders(tbl):
    b = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement('w:' + edge)
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '6')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        b.append(el)
    tbl._tbl.tblPr.append(b)


def cellpara(cell, text, bold=False, italic=False, size=14, red=False,
             align=AL.CENTER, first=False, spacing=1.15):
    p = cell.paragraphs[0] if first else cell.add_paragraph()
    p.paragraph_format.alignment = align
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = spacing
    if text:
        add_run(p, text, bold, italic, size, red)
    return p


# ---------------- header ----------------
h = doc.add_table(rows=1, cols=2)
h.alignment = WD_TABLE_ALIGNMENT.CENTER
set_fixed(h, [Cm(5.5), Cm(10.5)])
L, R = h.cell(0, 0), h.cell(0, 1)
cellpara(L, 'CÔNG TY CỔ PHẦN', bold=True, size=13, first=True)
cellpara(L, 'THỊNH ĐẠT', bold=True, size=13)
cellpara(L, 'Số:         /ĐN-TĐ', size=13)
cellpara(R, 'CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM', bold=True, size=13, first=True)
cellpara(R, 'Độc lập – Tự do – Hạnh phúc', bold=True, size=13)
cellpara(R, 'Lào Cai, ngày      tháng      năm 2026', italic=True, size=13)

para('', space_after=8)
para('GIẤY ĐỀ NGHỊ', bold=True, align=AL.CENTER, indent=Cm(0), size=16, space_after=4)
para('Cấp giấy phép sử dụng vật liệu nổ công nghiệp', bold=True, align=AL.CENTER,
     indent=Cm(0), size=14, space_after=10)
para('Kính gửi: Sở Công Thương tỉnh Lào Cai', align=AL.CENTER, indent=Cm(0),
     size=14, space_after=10)

# ---------------- body theo Mẫu số 04 ----------------
para('Tên tổ chức, doanh nghiệp: Công ty Cổ phần Thịnh Đạt.')
para('Nơi đặt trụ sở chính: Tổ 5, phường Trung Tâm, tỉnh Lào Cai.')
para('Điện thoại: 02163 879828.')
para('Giấy chứng nhận đăng ký doanh nghiệp số 5200280642 do Sở Tài chính tỉnh Lào Cai cấp '
     'thay đổi lần thứ 14 ngày 05 tháng 11 năm 2025.')
para('Họ tên, địa chỉ, số định danh cá nhân của người đại diện theo pháp luật: Nguyễn Anh Tuấn; '
     'Tổ 8, phường Trung Tâm, tỉnh Lào Cai; số định danh cá nhân 015085000988.')
para('Giấy chứng nhận đủ điều kiện về an ninh, trật tự số 01/GCN-CĐ2 do Phòng Cảnh sát quản lý '
     'hành chính về trật tự xã hội, Công an tỉnh Yên Bái cấp ngày 09 tháng 01 năm 2024.')
para('Lý do đề nghị cấp: Giấy phép sử dụng vật liệu nổ công nghiệp số 32/GP-SCT ngày 26 tháng 7 '
     'năm 2021 của Sở Công Thương tỉnh Yên Bái đã hết thời hạn.')
para('Chủng loại, số lượng vật liệu nổ công nghiệp sử dụng:', space_before=3, space_after=4)

rows = [
    ['STT', 'Tên vật liệu nổ công nghiệp', 'Đơn vị tính', 'Số lượng', 'Ghi chú'],
    ['1', 'Thuốc nổ các loại (thuốc nổ AĐ1; thuốc nổ nhũ tương dùng cho mỏ hầm lò, công trình '
          'ngầm không có khí nổ)', 'kg/năm', ('40.848',), ''],
    ['2', 'Kíp nổ các loại (kíp nổ điện vi sai; kíp nổ điện số 8)', 'cái/năm', ('84.572',), ''],
    ['3', 'Dây nổ, dây cháy chậm', 'm/năm', '0', 'Không sử dụng'],
]
widths = [Cm(1.3), Cm(7.0), Cm(2.2), Cm(2.3), Cm(3.2)]
t = doc.add_table(rows=len(rows), cols=5)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
set_borders(t)
set_fixed(t, widths)
for i, row in enumerate(rows):
    for j, val in enumerate(row):
        cell = t.cell(i, j)
        cell.width = widths[j]
        red = isinstance(val, tuple)
        txt = val[0] if red else val
        p = cell.paragraphs[0]
        p.paragraph_format.alignment = AL.LEFT if (i > 0 and j == 1) else AL.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        add_run(p, txt, bold=(i == 0), size=13, red=red)

para('', space_after=6)
para('Địa điểm sử dụng vật liệu nổ công nghiệp: mỏ chì - kẽm Trống Pá Sang, thôn Kháo Nhà, '
     'xã Tú Lệ, tỉnh Lào Cai.')
para('Thời hạn sử dụng vật liệu nổ công nghiệp: kể từ ngày ký đến hết ngày 20 tháng 01 năm 2030.')
mixed([('Họ tên, số định danh cá nhân của người đến liên hệ: ', False),
       ('【họ và tên; số định danh cá nhân】', True), ('.', False)])
para('Đề nghị Sở Công Thương tỉnh Lào Cai xem xét và cấp giấy phép sử dụng vật liệu nổ công '
     'nghiệp cho doanh nghiệp theo quy định tại Luật Quản lý, sử dụng vũ khí, vật liệu nổ và '
     'công cụ hỗ trợ ngày 29 tháng 6 năm 2024./.', space_before=3)

# ---------------- nơi nhận + chữ ký ----------------
para('', space_after=6)
sg = doc.add_table(rows=1, cols=2)
sg.alignment = WD_TABLE_ALIGNMENT.CENTER
set_fixed(sg, [Cm(7.0), Cm(9.0)])
SL, SR = sg.cell(0, 0), sg.cell(0, 1)
cellpara(SL, 'Nơi nhận:', bold=True, italic=True, size=12, align=AL.LEFT, first=True)
cellpara(SL, '- Như trên;', size=12, align=AL.LEFT)
cellpara(SL, '- Lưu: VT.', size=12, align=AL.LEFT)
cellpara(SR, 'GIÁM ĐỐC', bold=True, size=14, first=True)
cellpara(SR, '(Chữ ký, dấu)', italic=True, size=12)
for _ in range(3):
    cellpara(SR, '')
cellpara(SR, 'Nguyễn Anh Tuấn', bold=True, size=14)

out = './Giay-de-nghi-cap-GP-su-dung-VLNCN-Thinh-Dat.docx'
doc.save(out)
print('saved', out)
