# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RED = RGBColor(0xC0, 0x00, 0x00)
IND = Cm(1.0)

doc = Document()

# ---- page setup ----
s = doc.sections[0]
s.page_width, s.page_height = Cm(21.0), Cm(29.7)
s.left_margin, s.right_margin = Cm(3.0), Cm(2.0)
s.top_margin, s.bottom_margin = Cm(2.0), Cm(2.0)

st = doc.styles['Normal']
st.font.name = 'Times New Roman'
st.font.size = Pt(14)
st.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
pf = st.paragraph_format
pf.space_after = Pt(0)
pf.space_before = Pt(0)
pf.line_spacing = 1.3
pf.alignment = AL.JUSTIFY


def para(text='', bold=False, italic=False, align=AL.JUSTIFY, indent=IND,
         size=14, red=False, space_before=0, space_after=0, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = align
    p.paragraph_format.first_line_indent = indent
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if keep:
        p.paragraph_format.keep_with_next = True
    if text:
        add_run(p, text, bold, italic, size, red)
    return p


def add_run(p, text, bold=False, italic=False, size=14, red=False):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    if red:
        r.font.color.rgb = RED
    return r


def set_fixed(tbl, widths):
    tblPr = tbl._tbl.tblPr
    lay = OxmlElement('w:tblLayout')
    lay.set(qn('w:type'), 'fixed')
    tblPr.append(lay)
    grid = tbl._tbl.find(qn('w:tblGrid'))
    if grid is not None:
        tbl._tbl.remove(grid)
    grid = OxmlElement('w:tblGrid')
    for w in widths:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(int(w.twips)))
        grid.append(gc)
    tbl._tbl.insert(1, grid)


def set_borders(tbl):
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement('w:' + edge)
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '6')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    tblPr.append(borders)


def repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    el = OxmlElement('w:tblHeader')
    el.set(qn('w:val'), 'true')
    trPr.append(el)


def table(rows, widths=None, size=12, header=True, aligns=None):
    n = len(rows[0])
    tbl = doc.add_table(rows=len(rows), cols=n)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    set_borders(tbl)
    if widths:
        set_fixed(tbl, widths)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = tbl.cell(i, j)
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.0
            if i == 0 and header:
                p.paragraph_format.alignment = AL.CENTER
            elif aligns:
                p.paragraph_format.alignment = aligns[j]
            else:
                p.paragraph_format.alignment = AL.CENTER
            red = isinstance(val, tuple)
            txt = val[0] if red else val
            add_run(p, str(txt), bold=(i == 0 and header), size=size, red=red)
            if widths:
                cell.width = widths[j]
    if header:
        repeat_header(tbl.rows[0])
    if widths:
        for r in tbl.rows:
            for j, c in enumerate(r.cells):
                c.width = widths[j]
    return tbl



def new_section(landscape=False):
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    if landscape:
        sec.orientation = WD_ORIENT.LANDSCAPE
        sec.page_width, sec.page_height = Cm(29.7), Cm(21.0)
        sec.left_margin, sec.right_margin = Cm(2.0), Cm(2.0)
        sec.top_margin, sec.bottom_margin = Cm(2.0), Cm(2.0)
    else:
        sec.orientation = WD_ORIENT.PORTRAIT
        sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
        sec.left_margin, sec.right_margin = Cm(3.0), Cm(2.0)
        sec.top_margin, sec.bottom_margin = Cm(2.0), Cm(2.0)
    return sec


def caption(text):
    para(text, bold=True, align=AL.CENTER, indent=Cm(0), size=13,
         space_before=6, space_after=4, keep=True)


def spacer(pt=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pt)
    p.paragraph_format.line_spacing = 1.0
    return p


# ================= HEADER =================
h = doc.add_table(rows=1, cols=2)
h.alignment = WD_TABLE_ALIGNMENT.CENTER
c1, c2 = h.cell(0, 0), h.cell(0, 1)
set_fixed(h, [Cm(5.0), Cm(11.0)])
c1.width, c2.width = Cm(5.0), Cm(11.0)
p = c1.paragraphs[0]
p.paragraph_format.alignment = AL.CENTER
p.paragraph_format.first_line_indent = Cm(0)
p.paragraph_format.line_spacing = 1.1
add_run(p, 'CÔNG TY CỔ PHẦN\nTHỊNH ĐẠT', bold=True, size=13)
p = c2.paragraphs[0]
p.paragraph_format.alignment = AL.CENTER
p.paragraph_format.first_line_indent = Cm(0)
p.paragraph_format.line_spacing = 1.1
add_run(p, 'CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM', bold=True, size=13)
p2 = c2.add_paragraph()
p2.paragraph_format.alignment = AL.CENTER
p2.paragraph_format.first_line_indent = Cm(0)
add_run(p2, 'Độc lập – Tự do – Hạnh phúc', bold=True, size=13)

spacer(10)
para('PHƯƠNG ÁN NỔ MÌN', bold=True, align=AL.CENTER, indent=Cm(0), size=16,
     space_after=4)
para('Khai thác quặng chì - kẽm bằng phương pháp hầm lò tại mỏ chì - kẽm '
     'Trống Pá Sang, xã Cao Phạ, huyện Mù Cang Chải, tỉnh Yên Bái '
     '(nay là xã Tú Lệ, tỉnh Lào Cai)',
     bold=True, align=AL.CENTER, indent=Cm(0), size=13, space_after=10)

# ================= I =================
para('I. CĂN CỨ LẬP PHƯƠNG ÁN', bold=True, space_before=6, space_after=3)
para('1. Căn cứ pháp lý, tiêu chuẩn, quy chuẩn', bold=True, space_after=2)
for t in [
    '- Luật Quản lý, sử dụng vũ khí, vật liệu nổ và công cụ hỗ trợ số 42/2024/QH15 ngày 29/6/2024, được sửa đổi, bổ sung tại Luật số 118/2025/QH15 ngày 10/12/2025;',
    '- Luật Địa chất và khoáng sản số 54/2024/QH15 ngày 29/11/2024;',
    '- Nghị định số 181/2024/NĐ-CP ngày 31/12/2024 của Chính phủ quy định chi tiết một số điều và biện pháp thi hành Luật Quản lý, sử dụng vũ khí, vật liệu nổ và công cụ hỗ trợ về vật liệu nổ công nghiệp và tiền chất thuốc nổ;',
    '- Thông tư số 23/2024/TT-BCT ngày 07/11/2024 của Bộ trưởng Bộ Công Thương quy định về quản lý, sử dụng vật liệu nổ công nghiệp, tiền chất thuốc nổ thuộc thẩm quyền quản lý của Bộ Công Thương, được sửa đổi, bổ sung tại Thông tư số 38/2025/TT-BCT ngày 19/6/2025 và Thông tư số 26/2026/TT-BCT ngày 20/5/2026;',
    '- QCVN 01:2019/BCT - Quy chuẩn kỹ thuật quốc gia về an toàn trong sản xuất, thử nghiệm, nghiệm thu, bảo quản, vận chuyển, sử dụng, tiêu hủy vật liệu nổ công nghiệp và bảo quản tiền chất thuốc nổ (ban hành kèm theo Thông tư số 32/2019/TT-BCT ngày 21/11/2019);',
    '- QCVN 04:2017/BCT - Quy chuẩn kỹ thuật quốc gia về an toàn trong khai thác quặng hầm lò (ban hành kèm theo Thông tư số 31/2017/TT-BCT ngày 28/12/2017);',
    '- QCVN 05:2012/BCT - Quy chuẩn kỹ thuật quốc gia về thuốc nổ nhũ tương dùng cho mỏ hầm lò, công trình ngầm không có khí và bụi nổ và các quy chuẩn kỹ thuật quốc gia tương ứng với chủng loại vật liệu nổ công nghiệp được lựa chọn sử dụng;',
    '- Giấy phép khai thác khoáng sản số 1339/GP-UBND ngày 20/7/2017 của Ủy ban nhân dân tỉnh Yên Bái cấp cho Công ty Cổ phần Thịnh Đạt được khai thác quặng chì - kẽm bằng phương pháp hầm lò tại khu vực Trống Pá Sang, xã Cao Phạ, huyện Mù Cang Chải, tỉnh Yên Bái (nay là xã Tú Lệ, tỉnh Lào Cai);',
    '- Quyết định số 86/QĐ-UBND ngày 23/01/2026 về việc thu hồi đất, chuyển mục đích sử dụng đất, cho Công ty Cổ phần Thịnh Đạt thuê đất tại thôn Kháo Nhà, xã Tú Lệ để thực hiện dự án;',
    '- Thuyết minh Thiết kế bản vẽ thi công điều chỉnh Dự án đầu tư xây dựng công trình khai thác quặng chì kẽm tại mỏ chì kẽm Trống Pá Sang, xã Cao Phạ, huyện Mù Cang Chải, tỉnh Yên Bái (nay là xã Tú Lệ, tỉnh Lào Cai) do Công ty TNHH MTV Tư vấn Đầu tư Xây dựng Công nghiệp Mỏ Luyện kim lập; được Công ty Cổ phần Tư vấn Khoáng sản và Tài nguyên Môi trường thẩm tra tại Báo cáo kết quả thẩm tra thiết kế xây dựng ngày 18/5/2026; Tổ thẩm định của Chủ đầu tư thông báo kết quả thẩm định tại Thông báo số 01/2026/TTĐ ngày 24/6/2026 (sau đây gọi tắt là Thiết kế được duyệt);',
    '- Quyết định số 09/QĐ-Cty ngày 24/6/2026 của Giám đốc Công ty Cổ phần Thịnh Đạt về việc phê duyệt thiết kế bản vẽ thi công điều chỉnh;',
    '- Văn bản số 2093/SCT-KTATMT ngày 18/9/2017 của Sở Công Thương tỉnh Yên Bái về việc thông báo kết quả thẩm định Thiết kế bản vẽ thi công công trình khai thác khoáng sản chì kẽm mỏ chì kẽm Trống Pá Sang (thiết kế giai đoạn trước);',
    '- Quyết định số 04/QĐ-Cty ngày 20/3/2026 của Giám đốc Công ty Cổ phần Thịnh Đạt về việc bổ nhiệm Chỉ huy nổ mìn tại mỏ chì kẽm Trống Pá Sang, xã Tú Lệ, tỉnh Lào Cai; Danh sách cán bộ, công nhân liên quan đến vật liệu nổ công nghiệp mỏ chì - kẽm xã Tú Lệ, tỉnh Lào Cai ngày 20/7/2026 của Công ty Cổ phần Thịnh Đạt;',
    '- Biên bản kiểm tra nghiệm thu hệ thống phòng cháy, chữa cháy công trình kho chứa vật liệu nổ công nghiệp ngày 10/6/2011 của Phòng Cảnh sát phòng cháy chữa cháy và cứu nạn cứu hộ, Công an tỉnh Yên Bái;',
]:
    para(t)

para('2. Quy mô, công suất, chế độ làm việc', bold=True, space_before=6, space_after=2)
for t in [
    '- Diện tích khu vực khai thác: 2,25 ha; phương pháp khai thác: hầm lò; biên giới dưới sâu của mỏ được giới hạn tại mức +1780 m.',
    '- Công suất khai thác: 5.000 tấn quặng nguyên khai/năm (mục 3.2.1 Thiết kế được duyệt, theo Giấy phép khai thác khoáng sản số 1339/GP-UBND).',
    '- Trữ lượng khai thác còn lại tính đến ngày 31/12/2025: 22.505,4 tấn; tuổi thọ còn lại của dự án 4,5 năm.',
    '- Chế độ làm việc: 300 ngày/năm; 01 ca/ngày đêm; 08 giờ/ca (Bảng 3.1 Thiết kế được duyệt).',
    '- Năng suất khai thác của một lò chợ theo Bảng 8.3 Thiết kế được duyệt: 12,18 tấn/chu kỳ khấu; 8,38 tấn/ngày đêm; 209,5 tấn/tháng; 628,5 tấn/quý; 2.513,5 tấn/năm. Mỏ tổ chức khai thác đồng thời 02 lò chợ để bảo đảm công suất 5.000 tấn/năm, tương ứng khoảng 16,7 tấn/ngày đêm; 416,7 tấn/tháng; 1.250 tấn/quý.',
    '- Khối lượng đường lò chuẩn bị còn phải thực hiện tiếp theo Thiết kế được duyệt là 395 m (Bảng 5.1 Thiết kế được duyệt); công tác đào lò chuẩn bị được thực hiện đồng thời với công tác khai thác lò chợ.',
]:
    para(t)

para('3. Phương pháp mở vỉa, hệ thống khai thác và trình tự khai thác', bold=True,
     space_before=6, space_after=2)
para('Mỏ được mở vỉa bằng các đường lò bằng (lò xuyên vỉa, lò dọc vỉa) kết hợp lò thượng; '
     'không có giếng mỏ, sân ga, hầm trạm. Mỏ tổ chức khai thác đồng thời trên 03 mức: '
     'mức khai thác thứ nhất từ +1890 m lên đến đầu lộ vỉa ranh giới phía Tây; mức khai thác '
     'thứ hai từ +1840 m đến +1890 m; mức khai thác thứ ba từ +1780 m đến +1840 m.')
para('Hệ thống khai thác áp dụng là hệ thống khai thác buồng lưu quặng cho thân quặng có góc dốc '
     'lớn hơn 50 độ; khấu quặng bằng phương pháp khoan nổ mìn lỗ khoan nhỏ. Trình tự khai thác: '
     'huy động trước các khu vỉa có điều kiện địa chất chắc chắn và thuận lợi; trong một vỉa, mức trên '
     'khai thác trước, mức dưới khai thác sau; trong mỗi mức, các cột khai thác được tiến hành lần lượt '
     'từ trong ra ngoài mặt bằng cửa lò. Theo Thiết kế được duyệt, mỏ tổ chức khai thác đồng thời 02 lò chợ.')
para('Thiết bị chủ yếu phục vụ công tác khoan nổ mìn theo Bảng 8.9 Thiết kế được duyệt: 02 búa '
     'khoan khí nén cầm tay YT-24; 01 máy nén khí công suất 29,1 m³/phút; 02 quạt gió cục bộ '
     'YBT-52-2; 02 máy nổ mìn BKM-1/100. Vận tải trong lò bằng goòng 0,8 m³ và máng trượt.',
     space_before=3)
para('Nhân công: 10 người/ngày đêm làm việc trực tiếp tại buồng khai thác và buồng thu hồi '
     '(Bảng 8.3 Thiết kế được duyệt); nhân sự trực tiếp thực hiện công tác vật liệu nổ công '
     'nghiệp gồm 07 người, nêu tại mục V.3 Phương án này.')

para('4. Các từ, cụm từ viết tắt', bold=True, space_before=6, space_after=2)
for t in [
    '- VLNCN: vật liệu nổ công nghiệp;',
    '- QCVN: quy chuẩn kỹ thuật quốc gia; TCVN: tiêu chuẩn quốc gia;',
    '- Thiết kế được duyệt: hồ sơ Thiết kế bản vẽ thi công điều chỉnh nêu tại mục I.1 Phương án này, được phê duyệt tại Quyết định số 09/QĐ-Cty ngày 24/6/2026;',
    '- Công ty: Công ty Cổ phần Thịnh Đạt.',
]:
    para(t)

# ================= II =================
para('II. ĐẶC ĐIỂM KHU VỰC NỔ MÌN', bold=True, space_before=10, space_after=3)
para('1. Vị trí, giới hạn tọa độ và cao độ khu vực nổ mìn', bold=True, space_after=2)
para('Toàn bộ khu vực nổ mìn nằm trọn trong ranh giới khu vực khai thác đã được cấp phép tại '
     'Giấy phép khai thác khoáng sản số 1339/GP-UBND ngày 20/7/2017, diện tích 2,25 ha, thuộc '
     'thôn Kháo Nhà, xã Tú Lệ, tỉnh Lào Cai (trước đây là bản Trống Pá Sang, xã Cao Phạ, huyện '
     'Mù Cang Chải, tỉnh Yên Bái). Tọa độ các điểm khép góc như sau:')
caption('Bảng 1. Tọa độ ranh giới khu vực khai thác (VN 2000, KTT 104°45’, múi chiếu 3°)')
table([
    ['Tên điểm khép góc', 'X (m)', 'Y (m)', 'Diện tích (ha)'],
    ['1', '2.412.129', '442.393', '2,25'],
    ['2', '2.412.174', '442.586', '2,25'],
    ['3', '2.412.187', '442.856', '2,25'],
    ['4', '2.412.135', '442.864', '2,25'],
    ['5', '2.412.107', '442.393', '2,25'],
], widths=[Cm(4.5), Cm(3.8), Cm(3.8), Cm(3.9)])
para('Nguồn: Bảng 2.6 Thiết kế được duyệt. Bản đồ địa hình khu vực nổ mìn thể hiện ranh giới khai '
     'trường, vị trí các mặt bằng cửa lò, hệ thống đường lò và vị trí kho vật liệu nổ công nghiệp '
     'được lập trên cơ sở bản vẽ của Thiết kế được duyệt và đính kèm Phương án này.',
     italic=True, size=12, space_before=3)
para('Các mặt bằng cửa lò: mặt bằng cửa lò số 2 tại mức +1890 m (520 m²); mặt bằng cửa lò số 3 tại '
     'mức +1840 m (1.820 m²); mặt bằng cửa lò số 4 tại mức +1780 m (2.000 m²); mặt bằng cửa lò số 5 '
     'tại mức +1750 m (2.000 m²).', space_before=4)

para('2. Đặc điểm địa hình, khí hậu', bold=True, space_before=6, space_after=2)
para('Địa hình khu vực có dạng dải kéo dài theo phương Tây Bắc - Đông Nam, thuộc vùng núi cao và '
     'trung bình, độ cao từ 1.400 m đến 2.200 m, độ cao giảm dần từ Tây sang Đông. Khí hậu mang '
     'đặc tính miền núi Tây Bắc, chia thành hai mùa rõ rệt; mùa mưa từ tháng 4 đến tháng 10, có '
     'nguy cơ lũ quét, trượt lở đất.')

para('3. Dân cư, nhà ở và công trình cần bảo vệ', bold=True, space_before=6, space_after=2)
para('Dân cư trong vùng thưa thớt, chủ yếu là đồng bào dân tộc Thái, H’Mông, Kinh, sống tập trung '
     'thành các điểm dân cư nhỏ ven đường, thung lũng và sườn đồi, cách xa khu vực mỏ. Mỏ cách Ngã Ba '
     'Kim (Quốc lộ 32) khoảng 10 km về phía Đông Bắc; từ Quốc lộ 32 vào mỏ đi theo tuyến đường liên '
     'thôn bê tông dài khoảng 5 km và tuyến đường đá cấp phối dài khoảng 5 km do Công ty đầu tư.')
p = para('', indent=IND, space_before=3)
add_run(p, 'Thống kê nhà ở, công trình không thuộc sở hữu của Công ty, kể cả công trình ngầm, trong bán '
           'kính 1.000 m tính từ các mặt bằng cửa lò và từ hình chiếu bằng của các gương nổ mìn: ')
add_run(p, 'không có nhà ở, cơ sở khám bệnh, chữa bệnh, di tích lịch sử - văn hóa, khu bảo tồn '
           'thiên nhiên, công trình quốc phòng, an ninh hoặc công trình quan trọng khác của quốc gia. '
           'Kết quả thống kê nêu trên được xác định lại bằng đo đạc thực tế tại hiện trường có sự '
           'chứng kiến của cơ quan quản lý nhà nước và Ủy ban nhân dân xã Tú Lệ.', red=True)
para('Theo mục 2.4.2 Chỉ dẫn kỹ thuật của Thiết kế được duyệt, khoảng cách an toàn đối với khu dân cư '
     'khi nổ mìn là không nhỏ hơn 500 m; khoảng cách thực tế từ khu vực nổ mìn đến điểm dân cư gần nhất '
     'lớn hơn trị số này.', space_before=3)

para('4. Đặc điểm địa chất, tính chất cơ lý đất đá và phân loại mỏ về khí, bụi nổ', bold=True,
     space_before=6, space_after=2)
for t in [
    '- Hệ số kiên cố tính toán của đá và quặng f = 8 (Bảng 8.10 Thiết kế được duyệt); trọng lượng thể tích trung bình của quặng 3,2 tấn/m³; chiều dày trung bình thân quặng 0,95 m; góc dốc trung bình thân quặng 74 độ.',
    '- Thành phần đất đá gồm đá phiến thạch anh, đá phiến sét silic, đá phiến silic, cát kết, bột kết có chứa các thân quặng chì kẽm công nghiệp. Tầng đất phủ mỏng, chiều sâu phân bố từ 0,1 m đến 0,3 m, đất mềm rời dễ sập lở.',
    '- Kết quả phân tích 05 mẫu thể trọng tại Thiết kế được duyệt: độ ẩm tự nhiên W = 0,35 ÷ 1,02%; khối lượng thể tích tự nhiên 3,11 ÷ 3,18 g/cm³; khối lượng thể tích khô 3,08 ÷ 3,17 g/cm³; khối lượng riêng 3,20 ÷ 3,22 g/cm³; độ lỗ rỗng n = 1,2 ÷ 3,8%. Thiết kế được duyệt không xác định các chỉ tiêu cường độ kháng nén σn và kháng kéo σk; toàn bộ tính toán khoan nổ mìn sử dụng hệ số kiên cố f = 8 theo Bảng 8.10 Thiết kế được duyệt.',
    '- Tầng chứa nước khe nứt nghèo nước; lưu lượng các nguồn lộ rất nhỏ, từ 0,001 l/s đến 0,1 l/s, xuất lộ dạng thấm rỉ.',
    '- Căn cứ Điều 52 QCVN 04:2017/BCT và kết quả phân tích mẫu hóa nhóm quặng, kết quả đo nồng độ H₂S trong đường lò: mỏ chì - kẽm Trống Pá Sang thuộc loại mỏ KHÔNG nguy hiểm về khí và bụi nổ (nồng độ lưu huỳnh nhỏ hơn 12%, nồng độ H₂S nhỏ hơn 6,6 ppm), được phép làm việc theo chế độ mỏ công tác bình thường.',
]:
    para(t)

para('5. Hướng, trình tự nổ mìn và ảnh hưởng đến công trình xung quanh', bold=True,
     space_before=6, space_after=2)
para('Toàn bộ công tác nổ mìn của mỏ được thực hiện trong hầm lò, không có nổ mìn lộ thiên. Do đó '
     'không phát sinh đá văng ra ngoài phạm vi đường lò và không phát sinh sóng đập không khí lan '
     'truyền trên mặt đất. Yếu tố ảnh hưởng cần kiểm soát là chấn động do nổ mìn đối với các công '
     'trình trên mặt đất, độ ổn định của đường lò và khí độc sau nổ mìn; các nội dung này được tính '
     'toán, quy định tại mục III.6 và mục IV Phương án này.')
para('Trình tự khai thác theo Thiết kế được duyệt: năm 2026 khai thác phần quặng còn lại của lò '
     'chợ cột số 18 và lò chợ cột số 5; các năm 2027, 2028, 2029 lần lượt khai thác các lò chợ cột '
     'số 6, 7, 8, 9. Khi chuyển sang khu vực khai thác mới, điều kiện địa chất, chiều dày và góc dốc '
     'thân quặng có thể thay đổi; chỉ huy nổ mìn phải xem xét thực tế gương lò để điều chỉnh số lỗ '
     'khoan và lượng thuốc nạp trong hộ chiếu nổ mìn, bảo đảm không vượt quá quy mô một đợt nổ quy '
     'định tại mục III.4 Phương án này.', space_before=3)

# ================= III =================
para('III. TÍNH TOÁN, LỰA CHỌN CÁC THÔNG SỐ KHOAN NỔ MÌN', bold=True,
     space_before=10, space_after=3)
para('Các thông số khoan nổ mìn tại Phương án này được lấy thống nhất với mục 8.1.2, mục 8.2.4, '
     'Bảng 8.2, Bảng 8.3 và Bảng 8.10 của Thiết kế được duyệt.')

para('1. Thông số khoan nổ mìn đào lò chuẩn bị', bold=True, space_before=6, space_after=2)
para('Chỉ tiêu thuốc nổ đơn vị được xác định theo công thức của G. M. Pokrovski '
     'q = q₁ × f꜀ × v × e × k_d (kg/m³); tổng số lỗ khoan trên gương N_G = N_b + N_rf; lượng thuốc nổ '
     'cho một chu kỳ đào lò Q = q × S_đ × η × L. Kết quả tính toán, lựa chọn cho từng loại đường lò '
     'như sau:')
new_section(landscape=True)
caption('Bảng 2. Thông số cơ bản của hộ chiếu khoan nổ mìn đào chống lò')
table([
    ['TT', 'Tên thông số', 'Đơn vị',
     'Lò xuyên vỉa, dọc vỉa VT và TG', 'Lò thượng (α ≤ 45°)',
     'Lò thượng (α > 45°)', 'Lò dọc vỉa phân tầng', 'Phễu tháo quặng'],
    ['1', 'Hình dạng tiết diện lò', '-', 'Hình thang', 'Hình thang', 'Hình chữ nhật', 'Hình thang', 'Hình chữ nhật'],
    ['2', 'Diện tích đào của lò (Sđ)', 'm²', '5,20', '4,20', '4,20', '3,60', '4,00'],
    ['3', 'Chiều rộng đường lò (B)', 'm', '2,31', '2,01', '1,70', '1,85', '1,85'],
    ['4', 'Hệ số kiên cố của đá, quặng (f)', '-', '8', '8', '8', '8', '8'],
    ['5', 'Thuốc nổ sử dụng', '-', 'AĐ1', 'AĐ1', 'AĐ1', 'AĐ1', 'AĐ1'],
    ['6', 'Đường kính thỏi thuốc (dt)', 'mm', '32', '32', '32', '32', '32'],
    ['7', 'Đường kính lỗ khoan (d = dt + 4 ÷ 8 mm)', 'mm', '36', '36', '36', '36', '36'],
    ['8', 'Khả năng sinh công của thuốc nổ (Pe)', 'cm³', '360', '330', '320', '330', '330'],
    ['9', 'Mật độ thuốc nổ (Δ)', 'kg/m³', '1.100', '1.100', '1.100', '1.100', '1.100'],
    ['10', 'Chỉ tiêu thuốc nổ đơn vị (q)', 'kg/m³', '2,52', '3,05', '3,15', '3,30', '3,13'],
    ['11', 'Tiến độ một chu kỳ khoan nổ (Lo)', 'm', '0,80', '0,80', '1,00', '0,80', '1,00'],
    ['12', 'Tổng số lỗ khoan trên gương (NG), trong đó:', 'lỗ', '22', '21', '21', '20', '21'],
    ['13', '- Số lỗ khoan biên, nền (Nb)', 'lỗ', '18', '16', '16', '15', '16'],
    ['14', '- Số lỗ khoan phá (Nf)', 'lỗ', '3', '4', '4', '4', '1'],
    ['15', '- Số lỗ khoan tạo rạch (Nr)', 'lỗ', '1', '1', '1', '1', '4'],
    ['16', 'Chiều sâu lỗ khoan biên, lỗ phá (L)', 'm', '1,00', '1,00', '1,20', '1,00', '1,20'],
    ['17', 'Chiều sâu lỗ khoan tạo rạch', 'm', '1,20', '1,20', '1,40', '1,20', '1,40'],
    ['18', 'Khoảng cách giữa các lỗ mìn tạo biên (rb)', 'm', '0,40', '0,40', '0,40', '0,40', '0,40'],
    ['19', 'Số thỏi thuốc lỗ biên / phá / rạch', 'thỏi', '1,5 / 2,0 / 2,5', '1,5 / 2,0 / 2,5',
     '2,5 / 3,0 / 3,5', '2,0 / 2,0 / 2,5', '2,0 / 2,0 / 2,5'],
    ['20', 'Chiều dài bua lỗ biên / phá / rạch', 'm', '0,625 / 0,500 / 0,575', '0,625 / 0,500 / 0,575',
     '0,575 / 0,450 / 0,525', '0,500 / 0,500 / 0,575', '0,700 / 0,700 / 0,775'],
    ['21', 'Lượng thuốc nổ một chu kỳ đào lò (Q)', 'kg', '7,10', '6,90', '11,10', '8,10', '8,80'],
    ['22', 'Số kíp nổ một chu kỳ (mỗi lỗ 01 kíp)', 'cái', '22', '21', '21', '20', '21'],
    ['23', 'Lượng thuốc nổ cho 01 m lò', 'kg/m', '8,88', '8,63', '11,10', '10,13', '8,80'],
    ['24', 'Số kíp nổ cho 01 m lò', 'cái/m', '28', '26', '22', '25', '22'],
], widths=[Cm(1.2), Cm(7.4), Cm(1.8), Cm(3.4), Cm(2.9), Cm(2.9), Cm(3.0), Cm(3.1)], size=11,
    aligns=[AL.CENTER, AL.LEFT, AL.CENTER, AL.CENTER, AL.CENTER, AL.CENTER, AL.CENTER, AL.CENTER])
para('Nguồn: Bảng 8.10 Thiết kế được duyệt. Khối lượng một thỏi thuốc 0,2 kg; chiều dài một thỏi 0,25 m.',
     italic=True, size=12, space_before=3, indent=Cm(0))
new_section(landscape=False)
para('Quy cách khoan: lỗ khoan tạo rạch khoan nghiêng khoảng 80 độ so với mặt gương và khoan sâu hơn '
     'các lỗ khác từ 0,15 m đến 0,20 m để tạo phễu nổ đầu tiên; lỗ khoan phá khoan vuông góc với mặt '
     'gương; lỗ khoan biên khoan cách đường biên thiết kế 10 cm, nghiêng 85 độ so với mặt phẳng gương lò. '
     'Bua được chế tạo từ hỗn hợp sét và cát, nạp chặt; chiều dài bua không nhỏ hơn trị số ghi tại '
     'Bảng 2 Phương án này. Trình tự khởi nổ: nhóm lỗ tạo rạch nổ trước, tiếp đến nhóm lỗ phá, sau cùng '
     'là nhóm lỗ biên và lỗ nền.', space_before=4)
para('Cấu trúc cột thuốc: nạp liên tục các thỏi thuốc đường kính 32 mm, chiều dài 0,25 m, khối '
     'lượng 0,2 kg/thỏi, tính từ đáy lỗ khoan; mìn mồi đặt tại thỏi thuốc trong cùng, kíp nổ lắp '
     'ngược hướng ra miệng lỗ. Chiều dài cột thuốc bằng số thỏi nhân với 0,25 m; phần lỗ khoan còn '
     'lại nạp bua với chiều dài theo Bảng 2 Phương án này. Đối với gương lò chợ, mỗi lỗ khoan nạp '
     '0,4 kg thuốc nổ (02 thỏi) và 01 kíp nổ, phần còn lại nạp bua.', space_before=3)

para('2. Thông số khoan nổ mìn khai thác lò chợ', bold=True, space_before=6, space_after=2)
para('Chỉ tiêu thuốc nổ đơn vị khai thác được xác định theo công thức của GS. Protodiakonov; '
     'chi phí thuốc nổ một chu kỳ Q = L꜀ × m × r × q.')
caption('Bảng 3. Thông số cơ bản của hộ chiếu khoan nổ mìn khai thác lò chợ')
table([
    ['TT', 'Tên thông số', 'Đơn vị', 'Giá trị'],
    ['1', 'Chiều dày trung bình thân quặng (m)', 'm', '0,95'],
    ['2', 'Góc dốc trung bình thân quặng', 'độ', '74'],
    ['3', 'Chiều dài lò chợ (Lc)', 'm', '30'],
    ['4', 'Tiến độ khấu gương một chu kỳ (r)', 'm', '1,0'],
    ['5', 'Chiều rộng một đợt nổ mìn', 'm', '5,0'],
    ['6', 'Chỉ tiêu thuốc nổ đơn vị (q)', 'kg/m³', '1,55'],
    ['7', 'Số hàng lỗ khoan trên gương', 'hàng', '3'],
    ['8', 'Khoảng cách giữa các hàng lỗ khoan', 'm', '0,3'],
    ['9', 'Khoảng cách giữa các lỗ khoan trong một hàng', 'm', '0,7'],
    ['10', 'Số lỗ khoan trong một hàng', 'lỗ', '42'],
    ['11', 'Tổng số lỗ khoan một chu kỳ khấu (3 × 42)', 'lỗ', '126'],
    ['12', 'Lượng thuốc nạp một lỗ khoan (cả 3 hàng)', 'kg/lỗ', '0,4'],
    ['13', 'Lượng thuốc nổ một chu kỳ khấu (126 × 0,4)', 'kg', '50,4'],
    ['14', 'Số kíp nổ một chu kỳ khấu (mỗi lỗ 01 kíp)', 'cái', '126'],
    ['15', 'Số đợt nổ trong một chu kỳ khấu (30 m : 5 m)', 'đợt', '6'],
    ['16', 'Quy mô một đợt nổ (50,4 : 6)', 'kg', '8,4'],
    ['17', 'Số kíp nổ một đợt (126 : 6)', 'cái', '21'],
    ['18', 'Sản lượng quặng nổ mìn một chu kỳ khấu', 'tấn', '12,18'],
], widths=[Cm(1.2), Cm(9.0), Cm(2.4), Cm(3.4)], size=12,
    aligns=[AL.CENTER, AL.LEFT, AL.CENTER, AL.CENTER])
para('Nguồn: mục 8.1.2 và Bảng 8.2 Thiết kế được duyệt. Chiều rộng trụ bảo vệ trong buồng khai thác '
     'bảo đảm a ≥ 2W = 1,52 m; thiết kế bố trí trụ bảo vệ thượng cột rộng 3 m, trụ bảo vệ lò dọc vỉa '
     'rộng 5 m nên bảo đảm không bị phá hủy khi nổ mìn.', italic=True, size=12, space_before=3)

para('3. Phương pháp nổ mìn, chủng loại vật liệu nổ công nghiệp, phương tiện khởi nổ', bold=True,
     space_before=6, space_after=2)
para('- Phương pháp nổ mìn: nổ mìn trong các lỗ khoan nhỏ, điều khiển nổ bằng kíp nổ điện vi sai, '
     'khởi nổ bằng máy nổ mìn chuyên dụng cho hầm lò. Không áp dụng nổ mìn buồng, nổ mìn lỗ khoan '
     'lớn và nổ mìn đốt bằng dây cháy chậm. Do toàn bộ công tác nổ mìn được thực hiện trong hầm lò '
     'nên không đặt ra thông số chiều cao tầng và đường cản chân tầng; thay vào đó sử dụng chiều dài '
     'một bước đào (tiến độ một chu kỳ khoan nổ) nêu tại Bảng 2 và Bảng 3 Phương án này.')
for t in [
    '- Thuốc nổ: thuốc nổ AĐ1 và thuốc nổ nhũ tương dùng cho mỏ hầm lò, thuộc Danh mục vật liệu nổ công nghiệp được phép sản xuất, kinh doanh, sử dụng tại Việt Nam ban hành kèm theo Phụ lục I Thông tư số 23/2024/TT-BCT, bảo đảm là loại được phép sử dụng trong mỏ hầm lò, công trình ngầm không có khí và bụi nổ (mỏ đã được xác định thuộc loại không nguy hiểm về khí và bụi nổ tại mục II.4 Phương án này). Thông số dùng trong tính toán tại Thiết kế được duyệt là khả năng sinh công 320 ÷ 360 cm³, mật độ 1.100 kg/m³, thỏi thuốc đường kính 32 mm, khối lượng 0,2 kg, chiều dài 0,25 m; khi lựa chọn loại thuốc nổ có thông số khác các trị số này, Công ty tính lại chỉ tiêu thuốc nổ đơn vị và lượng thuốc nạp cho từng lỗ mìn, thể hiện trong hộ chiếu nổ mìn của đợt nổ tương ứng.',
    '- Phụ kiện nổ: kíp nổ điện vi sai và kíp nổ điện số 8 thuộc Danh mục nêu trên, bảo đảm là loại được phép sử dụng trong mỏ hầm lò, công trình ngầm không có khí và bụi nổ; áp dụng thống nhất cho cả gương đào lò và gương lò chợ.',
    '- Phương tiện khởi nổ: máy nổ mìn BKM-1/100 (02 chiếc) hoặc máy nổ mìn chuyên dụng cho hầm lò có đặc tính kỹ thuật tương đương, được kiểm định theo quy định (Bảng 8.9 Thiết kế được duyệt). Máy đo điện trở kíp và mạng nổ được kiểm định 01 lần/06 tháng.',
    '- Thiết bị khoan: búa khoan khí nén cầm tay YT-24 (02 chiếc) hoặc loại có đặc tính kỹ thuật tương đương; máy nén khí công suất 29,1 m³/phút.',
]:
    para(t)
p = para('- ', indent=IND)
add_run(p, 'Công ty KHÔNG sử dụng dây cháy chậm và phương pháp nổ mìn đốt tại mỏ. Lý do: thân quặng '
           'có góc dốc trung bình 74 độ, các lò thượng của mỏ được đào theo góc dốc thân quặng '
           '(lớn hơn 45 độ), thuộc trường hợp không được nổ mìn bằng dây cháy chậm ở lò đứng, lò '
           'nghiêng có độ dốc trên 30 độ theo quy định tại QCVN 01:2019/BCT; đồng bộ thiết bị nổ mìn '
           'tại Bảng 8.9 Thiết kế được duyệt cũng chỉ trang bị máy nổ mìn điện.', bold=False)
para('Trường hợp kết quả đo, phân tích khí mỏ xác định mỏ chuyển sang loại nguy hiểm về khí hoặc bụi '
     'nổ theo QCVN 04:2017/BCT thì phải chuyển sang sử dụng thuốc nổ an toàn và kíp nổ an toàn dùng '
     'cho hầm lò có khí và bụi nổ, đồng thời điều chỉnh Phương án này trước khi tiếp tục thực hiện.',
     space_before=3)

para('4. Quy mô một đợt nổ và khối lượng thuốc nổ tức thời lớn nhất', bold=True,
     space_before=6, space_after=2)
for t in [
    '- Đối với gương đào lò: quy mô một đợt nổ lớn nhất là 11,10 kg thuốc nổ và 21 kíp nổ (gương lò thượng có góc dốc lớn hơn 45 độ - trường hợp có lượng thuốc nổ một chu kỳ lớn nhất trong Bảng 2 Phương án này).',
    '- Đối với gương lò chợ: quy mô một đợt nổ là 8,4 kg thuốc nổ và 21 kíp nổ; trường hợp nổ toàn bộ chiều dài lò chợ trong một chu kỳ khấu thì lượng thuốc nổ tối đa là 50,4 kg và 126 kíp nổ.',
    '- Như vậy, quy mô một đợt nổ lớn nhất của mỏ là 50,4 kg thuốc nổ, 126 kíp nổ; trị số này được sử dụng để tính toán khoảng cách an toàn tại mục III.6 Phương án này (lấy theo hướng an toàn hơn).',
    '- Toàn bộ các đợt nổ được điều khiển bằng kíp nổ điện vi sai. Mỗi cấp vi sai bố trí không quá 21 lỗ khoan; khối lượng thuốc nổ đồng thời lớn nhất trong một cấp vi sai không vượt quá 8,4 kg đối với gương lò chợ và không vượt quá 11,10 kg đối với gương đào lò. Số cấp vi sai, số lỗ và khối lượng thuốc nổ của từng cấp được thể hiện cụ thể trong hộ chiếu nổ mìn của từng đợt nổ.',
]:
    para(t)

para('5. Khối lượng vật liệu nổ công nghiệp sử dụng', bold=True, space_before=6, space_after=2)
p = para('', indent=IND)
add_run(p, 'Khối lượng VLNCN sử dụng hằng năm của mỏ được lấy theo Bảng 8.12 Thiết kế được duyệt, '
           'gồm: khai thác 4.960 kg thuốc nổ và 12.400 kíp nổ; đào lò 32.175 kg thuốc nổ và 64.484 kíp nổ; '
           'hệ số dự trữ K = 1,1 (do các thân quặng tại mỏ là quặng nhiệt dịch thường xuyên biến động). '
           'Tổng khối lượng lớn nhất trong 01 năm:', red=True)
caption('Bảng 4. Khối lượng vật liệu nổ công nghiệp sử dụng')
table([
    ['TT', 'Chủng loại vật liệu nổ công nghiệp', 'Đơn vị', 'Hằng năm', 'Hằng quý', 'Hằng tháng'],
    ['1', 'Thuốc nổ các loại (thuốc nổ AĐ1; thuốc nổ nhũ tương dùng cho mỏ hầm lò)', 'kg',
     ('40.848',), ('10.212',), ('3.404',)],
    ['2', 'Kíp nổ các loại (kíp nổ điện vi sai; kíp nổ điện số 8)', 'cái',
     ('84.572',), ('21.143',), ('7.048',)],
    ['3', 'Dây nổ, dây cháy chậm', 'm', '0', '0', '0'],
], widths=[Cm(1.1), Cm(6.0), Cm(1.6), Cm(2.5), Cm(2.4), Cm(2.4)], size=12,
    aligns=[AL.CENTER, AL.LEFT, AL.CENTER, AL.CENTER, AL.CENTER, AL.CENTER])
para('Ghi chú: khối lượng hằng quý, hằng tháng được xác định bằng cách chia đều khối lượng hằng năm '
     'và làm tròn; khối lượng thực tế từng tháng được điều chỉnh theo tiến độ đào lò, khai thác nhưng '
     'không vượt quá khối lượng hằng năm nêu trên.', italic=True, size=12, space_before=3)
para('Kho VLNCN của Công ty được xây dựng trên cao trình +1900 m, phía Bắc khai trường, diện tích '
     'sử dụng 35,02 m², sức chứa 3.000 kg thuốc nổ và phụ kiện nổ tương ứng. Do nhu cầu thuốc nổ lớn nhất trong một '
     'tháng (3.404 kg) lớn hơn sức chứa cho phép của kho, Công ty tổ chức mua và nhập kho làm nhiều '
     'đợt trong tháng; số lượng VLNCN mỗi lần nhập kho và khối lượng tồn chứa tại mọi thời điểm không '
     'vượt quá sức chứa cho phép của kho là 3.000 kg.', space_before=4)

para('6. Tính toán khoảng cách an toàn', bold=True, space_before=6, space_after=2)
para('a) Khoảng cách an toàn về chấn động đối với công trình trên mặt đất', bold=True, italic=True,
     space_after=2)
para('Áp dụng công thức tại Phụ lục 7 QCVN 01:2019/BCT: R꜀ = K꜀ × α × ∛Q (m); trong đó K꜀ là hệ số '
     'phụ thuộc tính chất nền đất của công trình cần bảo vệ, α là hệ số phụ thuộc chỉ số tác động nổ, '
     'Q là khối lượng thuốc nổ của một đợt nổ (kg).')
para('Với Q = 50,4 kg (quy mô một đợt nổ lớn nhất của mỏ theo mục III.4), ∛Q = 3,69. Lấy K꜀ × α = 7 '
     '(công trình nhà cấp IV, nền đất), được R꜀ = 7 × 3,69 = 25,8 m; làm tròn R꜀ = 26 m.')
para('b) Khoảng cách an toàn về tác động của sóng đập không khí', bold=True, italic=True,
     space_before=4, space_after=2)
para('Công tác nổ mìn được thực hiện hoàn toàn trong hầm lò kín, sóng đập không khí không lan truyền '
     'ra ngoài trời nên không áp dụng đối với công trình trên mặt đất. Trong hầm lò, người phải rút về '
     'vị trí an toàn theo quy định tại mục IV.3 Phương án này.')
para('c) Khoảng cách an toàn do đá văng', bold=True, italic=True, space_before=4, space_after=2)
para('Không phát sinh đá văng ra ngoài phạm vi đường lò do toàn bộ đợt nổ được thực hiện trong hầm lò. '
     'Tại khu vực cửa lò, trong thời gian nổ mìn phải cấm người và phương tiện trong bán kính 50 m tính '
     'từ cửa lò của mức đang nổ mìn.')
para('d) Khoảng cách an toàn đối với người trong hầm lò', bold=True, italic=True,
     space_before=4, space_after=2)
para('Theo mục 2.4.2 Chỉ dẫn kỹ thuật của Thiết kế được duyệt, khoảng cách an toàn đối với người trong '
     'hầm lò khi nổ mìn là tối thiểu 150 m đối với đường lò thẳng và tối thiểu 100 m đối với đường lò '
     'cong. Các trị số này lớn hơn khoảng cách tối thiểu 50 m quy định tại QCVN 01:2019/BCT đối với '
     'nổ mìn ở lò chợ, do đó Phương án lấy theo trị số của Thiết kế được duyệt (theo hướng an toàn hơn).')
para('đ) Khoảng cách an toàn đối với thiết bị', bold=True, italic=True,
     space_before=4, space_after=2)
para('Trước mỗi đợt nổ, toàn bộ búa khoan, đường ống khí nén mềm, goòng, cáp điện di động và dụng '
     'cụ tại gương phải được tháo và di chuyển ra ngoài phạm vi tối thiểu 50 m tính từ gương nổ '
     'hoặc đưa về vị trí có vì chống bảo vệ. Quạt gió cục bộ và đường ống gió được giữ nguyên để '
     'phục vụ thông gió sau nổ mìn. Cắt điện khu vực gương nổ trước giờ khởi nổ theo quy định tại '
     'mục IV.4 Phương án này.')
para('e) Kết luận', bold=True, italic=True, space_before=4, space_after=2)
para('Khoảng cách an toàn lớn nhất phải bảo đảm đối với công trình trên mặt đất là R꜀ = 26 m. Theo mục '
     'II.3 Phương án này, trong bán kính 1.000 m tính từ các mặt bằng cửa lò và từ hình chiếu bằng của '
     'các gương nổ mìn không có nhà ở, công trình không thuộc sở hữu của Công ty; khoảng cách đến điểm '
     'dân cư gần nhất lớn hơn 500 m. Do đó khu vực nổ mìn của mỏ bảo đảm khoảng cách an toàn và không '
     'thuộc trường hợp phải được cơ quan có thẩm quyền phê duyệt Phương án nổ mìn và sự đồng ý bằng văn '
     'bản của Ủy ban nhân dân cấp tỉnh theo quy định tại điểm d khoản 2 Điều 38 Luật số 42/2024/QH15; '
     'Phương án này do người quản lý của Công ty phê duyệt theo quy định tại điểm d khoản 1 Điều 39 '
     'Luật số 42/2024/QH15.')

para('7. Thời gian nổ mìn', bold=True, space_before=6, space_after=2)
para('Toàn bộ công tác nổ mìn của mỏ được thực hiện trong hầm lò; theo mục II.3 Phương án này, '
     'trong bán kính 1.000 m tính từ các mặt bằng cửa lò và từ hình chiếu bằng của các gương nổ '
     'mìn không có nhà ở, công trình không thuộc sở hữu của Công ty. Do đó công tác nổ mìn không '
     'phát sinh đá văng, sóng đập không khí và tiếng nổ ảnh hưởng đến khu dân cư trên mặt đất; '
     'Phương án không ấn định khung giờ cấm nổ mìn.')
for t in [
    '- Nổ mìn được thực hiện theo chu kỳ sản xuất của từng gương, ngay sau khi kết thúc công tác khoan, nạp mìn, đấu nối mạng nổ và đã hoàn thành việc rút toàn bộ người ra vị trí an toàn, đặt đủ trạm gác theo mục IV.7 Phương án này; kể cả vào ban đêm khi mỏ tổ chức sản xuất ca đêm.',
    '- Thời điểm nổ mìn cụ thể của từng đợt do chỉ huy nổ mìn quyết định và được ghi trong hộ chiếu nổ mìn của đợt nổ tương ứng.',
    '- Theo Bảng 3.1 Thiết kế được duyệt, mỏ tổ chức làm việc 300 ngày/năm, 01 ca/ngày đêm, 08 giờ/ca. Khi thay đổi số ca làm việc trong ngày đêm, Công ty thông báo trước bằng văn bản đến Ủy ban nhân dân xã Tú Lệ và Sở Công Thương.',
    '- Không thực hiện nổ mìn khi hệ thống thông gió, chiếu sáng, thông tin liên lạc của mỏ không bảo đảm hoặc khi có mưa lớn, giông sét (mục IV.9 Phương án này).',
]:
    para(t)
para('Trước khi nổ mìn lần đầu, Công ty thông báo bằng văn bản đến Ủy ban nhân dân xã Tú Lệ và '
     'trưởng thôn Kháo Nhà về địa điểm, thời gian hoạt động nổ mìn và quy ước tín hiệu cảnh báo.')

para('8. Cung ứng, bảo quản và vận chuyển vật liệu nổ công nghiệp', bold=True,
     space_before=6, space_after=2)
for t in [
    '- Công ty mua VLNCN của tổ chức có Giấy phép kinh doanh VLNCN theo hợp đồng; đơn vị cung ứng vận chuyển và bàn giao tại kho VLNCN của Công ty tại thôn Kháo Nhà, xã Tú Lệ, tỉnh Lào Cai, có biên bản bàn giao kèm theo.',
    '- Kho VLNCN của Công ty là kho nổi, cố định trên mặt đất, diện tích sử dụng 35,02 m² (gian chứa thuốc nổ 24,47 m², gian chứa phụ kiện nổ 10,55 m²), sức chứa 3.000 kg thuốc nổ và phụ kiện nổ tương ứng theo kết luận tại Biên bản kiểm tra nghiệm thu hệ thống phòng cháy, chữa cháy công trình kho chứa VLNCN ngày 10/6/2011 của Phòng Cảnh sát phòng cháy chữa cháy và cứu nạn cứu hộ, Công an tỉnh Yên Bái. Kho được canh gác 24/24 giờ.',
    '- Việc xuất, nhập kho thực hiện theo QCVN 01:2019/BCT, có sổ sách theo dõi, kiểm đếm hằng ngày. VLNCN sử dụng không hết trong ca phải nhập lại kho ngay trong ngày, có xác nhận của chỉ huy nổ mìn và thủ kho.',
    '- Việc vận chuyển VLNCN từ kho đến nơi sử dụng trong hầm lò được thực hiện theo quy định tại mục IV.1 Phương án này và QCVN 01:2019/BCT. Trường hợp vận chuyển VLNCN trên đường bộ ngoài phạm vi mỏ, Công ty thực hiện theo Giấy phép vận chuyển VLNCN do cơ quan Công an có thẩm quyền cấp.',
]:
    para(t)

# ================= IV =================
para('IV. BIỆN PHÁP BẢO ĐẢM AN TOÀN KHI NỔ MÌN', bold=True, space_before=10, space_after=3)

para('1. An toàn khi bốc dỡ, bảo quản tạm và vận chuyển vật liệu nổ công nghiệp', bold=True,
     space_after=2)
for t in [
    '- Khu vực bốc dỡ phải có biển báo xác định giới hạn ngăn cách; người không có nhiệm vụ không được vào khu vực đã ngăn cách; có lực lượng bảo vệ theo quy định.',
    '- Vận chuyển từ kho đến nơi sử dụng trong hầm lò bằng ô tô, xe goòng chạy trên đường ray hoặc mang xách thủ công. Không vận chuyển thuốc nổ và phụ kiện nổ trong cùng một chuyến; không xếp các hòm, túi đựng VLNCN cao hơn thành toa xe goòng.',
    '- Thợ mìn, người phục vụ phải mang theo đèn ắc quy phòng nổ hoạt động tốt khi vận chuyển VLNCN trong hầm lò; không được để goòng có VLNCN tự trôi theo độ dốc.',
    '- Khi bảo quản tạm tại nơi làm việc, VLNCN phải để trong hòm, thùng chứa theo Phụ lục 10 QCVN 01:2019/BCT, đặt tại vị trí cao không bị ngập nước, cách gương lò lớn hơn 30 m, dưới sự quản lý trực tiếp của thợ mìn hoặc người bảo vệ; kíp nổ để cách ly với thuốc nổ; số lượng không vượt quá nhu cầu 01 ca.',
]:
    para(t)

para('2. An toàn khi khoan và nạp mìn', bold=True, space_before=6, space_after=2)
for t in [
    '- Trước khi khoan phải kiểm tra tình trạng nóc lò, hông lò, cạy hết đá om, đá treo để tạo không gian làm việc an toàn; kiểm tra hệ thống khí nén, cần khoan và các van an toàn; đo nồng độ khí cháy nổ và khí độc, chỉ được khoan khi nồng độ khí trong giới hạn cho phép theo QCVN 04:2017/BCT.',
    '- Quản đốc hoặc cán bộ trực ca trực tiếp xem xét toàn bộ gương lò, đánh dấu vị trí lỗ khoan theo hộ chiếu. Khoan lỗ hàng nóc trước, lỗ giữa và lỗ hàng nền sau. Nghiêm cấm khoan tiếp hoặc khoan mới vào lỗ mìn đã nạp thuốc.',
    '- Trước khi nạp mìn phải kiểm tra lại nồng độ khí tại gương; chỉ được nạp mìn khi nồng độ khí CH₄ không lớn hơn 1% và không có hiện tượng xì khí. Công tác nạp do 02 người thực hiện: một người nạp, một người chuẩn bị bua; nạp lỗ hàng nóc trước, lỗ giữa và lỗ nền sau.',
    '- Bua nạp chặt, chiều dài không nhỏ hơn trị số ghi tại Bảng 2 Phương án này. Hai đầu dây kíp luôn được đấu chập với nhau cho đến khi đấu vào mạng nổ.',
    '- Không nạp và nổ mìn trong gương lò có khoảng chưa chống lớn hơn quy định trong thiết kế chống lò hoặc khi vì chống ở gương đã bị hư hỏng.',
]:
    para(t)

para('3. An toàn khi nổ mìn trong hầm lò', bold=True, space_before=6, space_after=2)
for t in [
    'a) Trước khi bắt đầu nạp mìn, theo hiệu lệnh của thợ mìn, tất cả mọi người trong khu vực gương lò phải rút ra vị trí an toàn; vị trí an toàn phải được thông gió bình thường, tránh được đất đá văng và được chống đỡ chắc chắn.',
    'b) Khoảng cách an toàn đối với người khi nổ mìn trong hầm lò là tối thiểu 150 m đối với đường lò thẳng và tối thiểu 100 m đối với đường lò cong (mục III.6.d Phương án này).',
    'c) Đối với các lò thượng, lò nghiêng có độ dốc lớn hơn 30 độ, chỉ được nổ mìn bằng kíp nổ điện vi sai; việc khởi nổ tiến hành từ nơi an toàn ngoài đường lò dốc. Nghiêm cấm nổ mìn bằng dây cháy chậm tại các đường lò này.',
    'd) Kể từ khi 02 gương lò còn cách nhau 20 m, không được nổ mìn đồng thời từ 02 gương đối diện; trước khi nạp mìn, người không có nhiệm vụ phải rút ra khỏi cả 02 gương. Khi 02 gương còn cách nhau 7,0 m, chỉ được tiến hành công tác ở một gương và phải khoan lỗ khoan thăm dò sâu hơn lỗ khoan từ 1,0 m trở lên.',
    'đ) Khi nổ mìn ở gương của một trong hai đường lò đào song song cách nhau không lớn hơn 20 m, tất cả mọi người phải rút ra khỏi cả 02 gương đến vị trí an toàn. Chỉ được khởi nổ sau khi đã nhận được thông báo mọi người đã rút hết và đã đặt trạm gác bảo vệ.',
    'e) Không được nổ mìn khi trong khoảng 20 m kể từ vị trí nổ mìn đi ra ngoài còn đất đá chưa xúc hết, toa xe, đồ vật chiếm trên một phần ba tiết diện ngang của lò làm cản trở việc đi lại.',
    'g) Khi các đường lò còn cách điểm bục từ 7 m đến 10 m, phải khoan thăm dò trước gương bằng choòng khoan dài từ 2,5 m đến 3,0 m để đề phòng sự cố bục nước, bục bùn.',
]:
    para(t)

para('4. An toàn khi nổ mìn điện', bold=True, space_before=6, space_after=2)
for t in [
    '- Không được bảo quản, vận chuyển kíp điện gần các nguồn thu, phát sóng điện từ tần số radio với khoảng cách nhỏ hơn khoảng cách quy định tại Phụ lục 6 QCVN 01:2019/BCT.',
    '- Toàn bộ kíp điện trong một mạng nổ phải cùng loại và cùng một nhà sản xuất; phải đo điện trở từng kíp trước khi sử dụng để kiểm tra sự phù hợp với giới hạn quy định của nhà chế tạo; sau khi đo, hai đầu dây phải được đấu chập lại và giữ ở trạng thái đó cho đến khi đấu vào mạng nổ. Dụng cụ đo có dòng điện phát vào mạch đo không vượt quá 50 mA, được kiểm định 01 lần/06 tháng.',
    '- Mạng điện nổ mìn luôn phải có hai dây dẫn; cấm sử dụng nước, đất, đường ống kim loại, đường ray, dây cáp để làm một trong hai dây dẫn. Đầu dây nối mạng phải được cạo sạch, mối nối chặt và quấn băng cách điện.',
    '- Chỉ được đấu nối mạng nổ sau khi đã nạp và lấp bua xong toàn bộ các phát mìn của một đợt nổ và những người không liên quan đã rút ra nơi an toàn. Cấm đấu mạng điện nổ mìn theo hướng đi từ nguồn điện đến các phát mìn.',
    '- Chìa khóa máy nổ mìn do chỉ huy nổ mìn giữ trong suốt thời gian từ khi chuẩn bị nạp đến khi khởi nổ; cấm giao chìa khóa cho người khác.',
    '- Cường độ dòng điện gây nổ phóng vào mỗi kíp không nhỏ hơn 1,0 A; không nhỏ hơn 1,3 A khi số kíp nổ đồng thời đến 300 chiếc và không nhỏ hơn 2,5 A khi khởi nổ bằng dòng điện xoay chiều.',
    '- Khi đóng cầu dao hoặc quay chìa khóa máy nổ mìn đến vị trí khởi nổ mà phát mìn không nổ, người khởi nổ phải tháo hai đầu dây dẫn chính ra khỏi nguồn điện, đấu chập lại và chỉ được vào kiểm tra bãi mìn sau ít nhất 05 phút.',
    '- Chỉ những thợ mìn đã qua đào tạo, huấn luyện và có ít nhất 01 năm kinh nghiệm làm việc với phương pháp nổ mìn điện mới được đấu, lắp mạng điện nổ mìn.',
]:
    para(t)

para('5. Thông gió và kiểm soát khí mỏ', bold=True, space_before=6, space_after=2)
for t in [
    '- Lượng không khí sạch đưa vào mỗi gương lò có nổ mìn phải bảo đảm để sau khi thông gió không quá 30 phút, hàm lượng khí độc sinh ra do nổ mìn tại đường lò người đi vào gương không vượt quá giới hạn cho phép theo QCVN 04:2017/BCT.',
    '- Sau khi nổ mìn phải thông gió tích cực bằng quạt cục bộ tối thiểu 30 phút; chỉ được đưa người vào làm việc sau khi đã đo kiểm tra hàm lượng ôxy, cacbonic, cacbon oxit và nhiệt độ bảo đảm quy định.',
    '- Bố trí cán bộ đo khí kiểm tra đầu ca và trong ca, ghi kết quả vào sổ theo dõi; không bố trí người vào làm việc tại những vị trí chưa được kiểm tra đo khí.',
]:
    para(t)

para('6. Biện pháp che chắn bảo vệ chống đá văng', bold=True, space_before=6, space_after=2)
para('Toàn bộ công tác nổ mìn được thực hiện trong hầm lò, đá văng bị giới hạn trong phạm vi đường '
     'lò nên không áp dụng biện pháp che chắn miệng lỗ mìn bằng lưới thép và bao cát như đối với nổ '
     'mìn lộ thiên. Biện pháp bảo vệ tương ứng là: rút toàn bộ người về vị trí an toàn có vì chống '
     'chắc chắn theo mục IV.3 Phương án này; cấm người và phương tiện trong bán kính 50 m tính từ '
     'cửa lò của mức đang nổ mìn; đóng kín cửa gió, cửa chắn tại các ngã ba, ngã tư đường lò dẫn vào '
     'khu vực nổ mìn.')

para('7. Tín hiệu cảnh báo và canh gác', bold=True, space_before=6, space_after=2)
for t in [
    '- Quy ước 03 hiệu lệnh: hiệu lệnh bắt đầu nạp mìn (01 hồi còi dài), hiệu lệnh khởi nổ (02 hồi còi ngắn), hiệu lệnh báo yên (03 hồi còi ngắn). Hiệu lệnh được phổ biến cho toàn bộ người lao động và niêm yết tại các mặt bằng cửa lò.',
    '- Bố trí trạm gác tại toàn bộ các ngã ba, ngã tư đường lò dẫn vào khu vực nổ mìn và tại cửa lò của mức đang nổ mìn. Mỗi trạm bố trí ít nhất 01 người, có băng gác, đèn và biển báo “Cấm vào”. Vị trí và số lượng trạm gác của từng đợt nổ được ghi cụ thể trong hộ chiếu nổ mìn.',
    '- Trong thời gian nổ mìn, cấm người và phương tiện trong bán kính 50 m tính từ cửa lò của mức đang nổ mìn.',
    '- Chỉ được bỏ trạm gác khi có lệnh của chỉ huy nổ mìn sau khi đã kiểm tra bãi mìn bảo đảm an toàn.',
]:
    para(t)

para('8. Kiểm tra sau khi nổ mìn và xử lý mìn câm', bold=True, space_before=6, space_after=2)
for t in [
    '- Sau khi nổ mìn và thông gió theo mục IV.5, cán bộ trực ca cùng thợ mìn vào kiểm tra chất lượng nổ mìn và tình trạng gương lò; cạy hết đá om, đá treo trước khi cho người vào làm việc và trước khi chống lò.',
    '- Thợ mìn phải đếm số phát mìn đã nổ. Trường hợp không đếm được hoặc có phát mìn không nổ, chỉ được trở lại khu vực bãi mìn sau 15 phút kể từ khi phát mìn cuối cùng nổ và sau khi đã thông gió hết khói mìn.',
    '- Khi phát hiện mìn câm, thợ mìn phải báo ngay cho chỉ huy nổ mìn; cắm biển cảnh báo, giữ nguyên trạm gác và chỉ được xử lý theo chỉ đạo trực tiếp của chỉ huy nổ mìn: khoan lỗ song song cách lỗ mìn câm không nhỏ hơn 30 cm để nổ phá; tuyệt đối không được khoan trực tiếp vào lỗ mìn câm, không được dùng dụng cụ kim loại moi lấy thuốc nổ. Chỉ cho người vào làm việc sau khi đã xử lý xong và được chỉ huy nổ mìn xác nhận an toàn.',
    '- Mọi trường hợp mìn câm phải được ghi vào hộ chiếu nổ mìn và sổ theo dõi; báo cáo Sở Công Thương theo quy định.',
]:
    para(t)

para('9. Ứng phó sự cố, thời tiết bất lợi và bảo đảm an ninh, trật tự', bold=True,
     space_before=6, space_after=2)
for t in [
    '- Không thực hiện nổ mìn khi có mưa lớn, giông sét hoặc khi hệ thống thông gió, chiếu sáng, thông tin liên lạc của mỏ không bảo đảm.',
    '- Công ty lập Bản đánh giá nguy cơ rủi ro về an toàn theo Điều 14 Thông tư số 23/2024/TT-BCT và rà soát hằng năm; xây dựng, phê duyệt và tổ chức luyện tập Kế hoạch ứng cứu khẩn cấp theo Điều 16 Thông tư số 23/2024/TT-BCT.',
    '- Tăng cường bảo vệ kho VLNCN và khu vực nổ mìn; trang bị camera quan sát 24/24 giờ tại cửa ra vào nhà kho, cổng kho và các vị trí xung yếu; kiểm đếm VLNCN xuất, nhập, tồn hằng ngày; báo cáo ngay cơ quan Công an và Sở Công Thương khi phát hiện mất mát VLNCN.',
]:
    para(t)

# ================= V =================
para('V. TỔ CHỨC THỰC HIỆN', bold=True, space_before=10, space_after=3)
para('1. Mỗi đợt nổ mìn phải lập hộ chiếu nổ mìn theo Mẫu số 02 Phụ lục VIII Thông tư số '
     '23/2024/TT-BCT trên cơ sở Phương án này, do chỉ huy nổ mìn lập và phê duyệt; nghiêm cấm thực '
     'hiện nổ mìn khi chưa có hộ chiếu nổ mìn. Trình tự và thủ tục kiểm soát các bước khoan, nạp, đấu '
     'nối mạng nổ, khởi nổ, canh gác và xử lý sau nổ mìn được thực hiện dưới sự giám sát trực tiếp của '
     'chỉ huy nổ mìn.')
para('2. Trách nhiệm cụ thể:', space_before=4)
for t in [
    '- Người quản lý về VLNCN: tổ chức thực hiện, kiểm tra, giám sát toàn bộ công tác quản lý, bảo quản, vận chuyển và sử dụng VLNCN; trình Giám đốc Công ty phê duyệt và giám sát việc thực hiện Phương án này.',
    '- Chỉ huy nổ mìn: lập hộ chiếu nổ mìn; trực tiếp chỉ huy các khâu nạp, đấu nối mạng nổ, khởi nổ, kiểm tra sau nổ mìn và xử lý mìn câm; giữ chìa khóa máy nổ mìn.',
    '- Thủ kho: quản lý xuất, nhập, tồn kho; lập và lưu sổ sách theo QCVN 01:2019/BCT.',
    '- Thợ nổ mìn, người phục vụ, người gác mìn: thực hiện đúng nhiệm vụ được phân công theo hộ chiếu nổ mìn.',
    '- Chỉ những người có Giấy chứng nhận huấn luyện kỹ thuật an toàn VLNCN còn hiệu lực mới được tham gia các công việc liên quan đến VLNCN; người quản lý, chỉ huy nổ mìn, thợ mìn, thủ kho phải bảo đảm trình độ chuyên môn theo Điều 4 Nghị định số 181/2024/NĐ-CP.',
]:
    para(t)
para('3. Nhân sự trực tiếp thực hiện Phương án này (theo Danh sách cán bộ, công nhân liên quan '
     'đến vật liệu nổ công nghiệp ngày 20/7/2026 của Công ty):', space_before=4)
for t in [
    '- Người quản lý về VLNCN kiêm Giám đốc điều hành mỏ: ông Bùi Đức Long, sinh năm 1985, trình độ Đại học Mỏ - Địa chất, Giấy chứng nhận huấn luyện kỹ thuật an toàn VLNCN số 28/SCT-2025.',
    '- Chỉ huy nổ mìn kiêm người phụ trách an toàn: ông Trần Mạnh Khải, sinh ngày 05/11/1991, được bổ nhiệm tại Quyết định số 04/QĐ-Cty ngày 20/3/2026 của Giám đốc Công ty, Giấy chứng nhận huấn luyện kỹ thuật an toàn VLNCN số 64/SCT-2025.',
    '- Thợ mìn: ông Đặng Dùn Và (Giấy chứng nhận số 29/SCT-2025); ông Hảng A Vàng (Giấy chứng nhận số 393/GCN-SCT).',
    '- Thủ kho VLNCN: ông Nguyễn Văn Toản (Giấy chứng nhận số 60/SCT-2024).',
    '- Người phục vụ, bảo vệ: ông Lý A Dì (Giấy chứng nhận số 26/SCT-2025); ông Nguyễn Văn Hải (Giấy chứng nhận số 30/SCT-2025).',
]:
    para(t)
para('4. Mọi sự kiện bất thường trong đợt nổ mìn, kể cả khi chưa đến mức xảy ra sự cố, phải được ghi '
     'vào hộ chiếu nổ mìn và sổ theo dõi; báo cáo cơ quan quản lý nhà nước có thẩm quyền theo Điều 17 '
     'Thông tư số 23/2024/TT-BCT.', space_before=4)
para('5. Công ty tổ chức giám sát ảnh hưởng của nổ mìn đến công trình, nhà ở của tổ chức, cá nhân xung '
     'quanh; tiếp nhận, giải quyết phản ánh và bồi thường thiệt hại (nếu có) theo quy định của pháp luật.')
para('6. Cá nhân, bộ phận vi phạm Phương án này bị xử lý theo nội quy lao động của Công ty và quy định '
     'của pháp luật.')
para('7. Phương án này có hiệu lực kể từ ngày được Giám đốc Công ty phê duyệt và Công ty được cấp Giấy '
     'phép sử dụng vật liệu nổ công nghiệp.')
para('8. Người lập Phương án: ông Trần Mạnh Khải, Chỉ huy nổ mìn của Công ty. Người phê duyệt '
     'Phương án: ông Nguyễn Anh Tuấn, Giám đốc Công ty Cổ phần Thịnh Đạt. Phương án không thuộc '
     'trường hợp phải được cơ quan có thẩm quyền phê duyệt và được sự đồng ý bằng văn bản của Ủy ban '
     'nhân dân cấp tỉnh theo điểm d khoản 2 Điều 38 Luật số 42/2024/QH15; lý do nêu tại mục III.6.e '
     'Phương án này.')
para('9. Khi có thay đổi về thiết kế, quy mô khai thác, chủng loại vật liệu nổ công nghiệp, điều kiện '
     'an toàn hoặc kết quả đo đạc khoảng cách thực tế đến nhà ở, công trình xung quanh, Phương án phải '
     'được rà soát, điều chỉnh và phê duyệt lại trước khi tiếp tục thực hiện.')
para('Trên đây là Phương án nổ mìn khai thác quặng chì - kẽm bằng phương pháp hầm lò tại mỏ chì - kẽm '
     'Trống Pá Sang, xã Tú Lệ, tỉnh Lào Cai của Công ty Cổ phần Thịnh Đạt. Công ty cam kết tổ chức '
     'thực hiện đúng các nội dung của Phương án và các quy định của pháp luật về quản lý, sử dụng vật '
     'liệu nổ công nghiệp./.', space_before=6)

# ---- signature ----
spacer(10)
sg = doc.add_table(rows=1, cols=2)
sg.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, (l1, l2) in enumerate([
    ('PHÊ DUYỆT\nGIÁM ĐỐC CÔNG TY', 'Nguyễn Anh Tuấn'),
    ('NGƯỜI LẬP PHƯƠNG ÁN\nCHỈ HUY NỔ MÌN', 'Trần Mạnh Khải'),
]):
    cell = sg.cell(0, j)
    cell.width = Cm(8.0)
    p = cell.paragraphs[0]
    p.paragraph_format.alignment = AL.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.15
    add_run(p, l1, bold=True, size=13)
    for _ in range(4):
        q = cell.add_paragraph()
        q.paragraph_format.first_line_indent = Cm(0)
    q = cell.add_paragraph()
    q.paragraph_format.alignment = AL.CENTER
    q.paragraph_format.first_line_indent = Cm(0)
    add_run(q, l2, bold=True, size=13)

out = './Phuong-an-no-min-Thinh-Dat-mo-chi-kem-Trong-Pa-Sang.docx'
doc.save(out)
print('saved', out)
