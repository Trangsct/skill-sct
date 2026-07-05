#!/usr/bin/env python3
"""
fill_template.py — Thư viện sửa nội dung file .docx mẫu mà GIỮ NGUYÊN
mọi format (khung, font, gạch chân, bảng chữ ký...).

Cách hoạt động:
  - Mở file mẫu .docx (đã được trau chuốt sẵn)
  - Sửa text trong các paragraph cụ thể
  - KHÔNG động đến cấu trúc table, không thêm/xóa run
  - Lưu thành file mới

Sử dụng:
  from fill_template import TemplateDoc
  doc = TemplateDoc('templates/01-cong-van.docx')
  doc.replace_in_table(table_idx=0, cell_idx=(0,0), pattern='Số:       /SCT-CN', replacement='Số: 123/SCT-CN')
  doc.replace_in_paragraph(idx=2, pattern='…………..', replacement='Ủy ban nhân dân tỉnh Lào Cai')
  doc.save('output/cv-moi.docx')
"""

from docx import Document
from copy import deepcopy
import re


class TemplateDoc:
    def __init__(self, template_path):
        self.template_path = template_path
        self.doc = Document(template_path)

    # ============================================================
    # TIỆN ÍCH NỘI BỘ
    # ============================================================

    @staticmethod
    def _is_shape_run(run):
        """Kiểm tra run có chứa shape (drawing/picture/altcontent) không."""
        xml = run._element.xml
        return ('<w:drawing' in xml
                or 'AlternateContent' in xml
                or '<w:pict' in xml
                or '<mc:AlternateContent' in xml)

    @classmethod
    def _get_text_runs(cls, paragraph):
        """Trả về danh sách runs CHỨA TEXT (không phải shape)."""
        return [run for run in paragraph.runs if not cls._is_shape_run(run)]

    @staticmethod
    def _norm_inline(text):
        """
        GUARD CHỐNG LỖI XUỐNG DÒNG Ở TIÊU ĐỀ/HEADER (lỗi tái diễn).

        Trong VBHC, một paragraph KHÔNG bao giờ chứa ngắt dòng cứng (<w:br/>):
        mọi việc xuống dòng đều phải tách thành paragraph riêng. python-docx tự
        chuyển '\\n'/'\\r' trong run.text thành <w:br/>, gây ngắt dòng xấu giữa
        cụm từ ở V/v, tên cơ quan, ngày tháng (vd 'Kế / hoạch phát / triển...').

        Hàm này thay mọi '\\r\\n', '\\r', '\\n' bằng MỘT dấu cách, đồng thời gộp
        dấu cách thừa do việc thay thế sinh ra ở ranh giới ngắt dòng — nhưng
        KHÔNG đụng tới các chuỗi nhiều dấu cách căn chỉnh có chủ đích trong
        header (vd 'Số:        /SCT-CN', 'ngày      tháng 6 năm 2026').
        """
        if text is None:
            return text
        # Chỉ thay newline thành 1 dấu cách; gộp khoảng trắng PHÁT SINH quanh
        # vị trí newline để tránh '  ' do nối 'phát\\ntriển' → 'phát triển'.
        import re as _re
        return _re.sub(r'[ \t]*[\r\n]+[ \t]*', ' ', text)

    @staticmethod
    def _dominant_run(text_runs):
        """
        Chọn run mang ĐỊNH DẠNG CHỦ ĐẠO của paragraph = text-run có text
        dài nhất (hòa thì lấy run đứng trước).

        Vì sao (bài học SZ13 05/7/2026): dòng ngày ở header thường bị Word
        tách thành nhiều run — run đầu rỗng/ngắn KHÔNG có w:sz tường minh,
        run dài phía sau mới mang w:sz=26 (13pt) + w:i (nghiêng). Nếu gán
        text mới vào run[0] rồi xóa run sau thì mất 13pt tường minh →
        qa_pdf_check báo [FAIL SZ13]. Gán vào run chủ đạo giữ được định dạng
        thật; các run khác xóa trắng nên thứ tự hiển thị không đổi.
        """
        return max(text_runs, key=lambda r: len(r.text))

    @classmethod
    def _replace_text_in_paragraph(cls, paragraph, pattern, replacement):
        """
        Sửa text trong paragraph, GIỮ NGUYÊN các run chứa shape (drawing).
        - Ghép text từ các text-run (bỏ qua shape-run)
        - Replace pattern → text mới
        - Gán new_text vào text-run CHỦ ĐẠO (text dài nhất — giữ đúng
          định dạng tường minh w:sz/w:i, xem _dominant_run)
        - Xóa text các text-run còn lại (KHÔNG động vào shape-run)
        """
        text_runs = cls._get_text_runs(paragraph)
        if not text_runs:
            return False
        full_text = ''.join(run.text for run in text_runs)
        if pattern not in full_text:
            return False
        new_text = cls._norm_inline(full_text.replace(pattern, replacement))
        target = cls._dominant_run(text_runs)
        for run in text_runs:
            run.text = new_text if run is target else ''
        return True

    @classmethod
    def _set_paragraph_text(cls, paragraph, new_text):
        """
        Đặt nội dung paragraph thành new_text, GIỮ NGUYÊN shape-runs.
        Gán vào run CHỦ ĐẠO (text dài nhất) để giữ định dạng tường minh
        (w:sz/w:i) — xem _dominant_run.
        """
        new_text = cls._norm_inline(new_text)
        text_runs = cls._get_text_runs(paragraph)
        if not text_runs:
            paragraph.add_run(new_text)
            return
        target = cls._dominant_run(text_runs)
        for run in text_runs:
            run.text = new_text if run is target else ''

    # ============================================================
    # API SỬA NỘI DUNG
    # ============================================================

    def replace_in_paragraph(self, idx, pattern, replacement):
        """Sửa text trong paragraph thứ idx (đếm từ 0)."""
        if idx >= len(self.doc.paragraphs):
            raise IndexError(f'Không có paragraph {idx}')
        return self._replace_text_in_paragraph(self.doc.paragraphs[idx], pattern, replacement)

    def set_paragraph_text(self, idx, new_text):
        """Đặt toàn bộ nội dung paragraph thứ idx."""
        if idx >= len(self.doc.paragraphs):
            raise IndexError(f'Không có paragraph {idx}')
        self._set_paragraph_text(self.doc.paragraphs[idx], new_text)

    def replace_in_cell(self, table_idx, row, col, pattern, replacement):
        """Sửa text trong ô bảng (table[table_idx].rows[row].cells[col])."""
        cell = self.doc.tables[table_idx].rows[row].cells[col]
        for p in cell.paragraphs:
            self._replace_text_in_paragraph(p, pattern, replacement)

    def replace_in_cell_paragraph(self, table_idx, row, col, para_idx, pattern, replacement):
        """Sửa text trong một paragraph cụ thể trong ô bảng."""
        cell = self.doc.tables[table_idx].rows[row].cells[col]
        if para_idx >= len(cell.paragraphs):
            return False
        return self._replace_text_in_paragraph(cell.paragraphs[para_idx], pattern, replacement)

    def set_cell_paragraph_text(self, table_idx, row, col, para_idx, new_text):
        """Đặt nội dung paragraph thứ para_idx trong ô bảng."""
        cell = self.doc.tables[table_idx].rows[row].cells[col]
        if para_idx >= len(cell.paragraphs):
            raise IndexError(f'Cell ({row},{col}) không có paragraph {para_idx}')
        self._set_paragraph_text(cell.paragraphs[para_idx], new_text)

    def replace_keeping_first_run(self, p_idx, new_text_after, separator=': '):
        """
        Sửa nội dung paragraph có "tiền tố bold" + "phần thân không bold".
        Ví dụ: "Điều 1: Thành lập..." → text-run[0]="Điều 1" bold, các text-run sau là nội dung.
        - Giữ NGUYÊN text-run[0] (đã bold sẵn) và các shape-run
        - Đặt text-run[1].text = separator + new_text_after
        - Xóa text các text-run còn lại
        """
        p = self.doc.paragraphs[p_idx]
        text_runs = self._get_text_runs(p)
        if len(text_runs) < 2:
            return False
        text_runs[1].text = separator + new_text_after
        for run in text_runs[2:]:
            run.text = ''
        return True

    # ============================================================
    # API SỬA NHIỀU PLACEHOLDER (find & replace toàn doc)
    # ============================================================

    def replace_all(self, pattern, replacement):
        """Tìm pattern trong toàn document (paragraph + cell) và replace."""
        # Trong paragraphs
        for p in self.doc.paragraphs:
            self._replace_text_in_paragraph(p, pattern, replacement)
        # Trong tables
        for t in self.doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        self._replace_text_in_paragraph(p, pattern, replacement)

    # ============================================================
    # API SỬA PHẦN THÂN (XÓA + INSERT)
    # ============================================================

    def replace_body_paragraphs(self, start_idx, end_idx, new_paragraphs):
        """
        Thay các paragraph từ start_idx đến end_idx (không bao gồm end_idx)
        bằng danh sách new_paragraphs (list dict {text, bold, italic, indent}).
        Giữ nguyên format của paragraph đầu tiên (start_idx) làm template.

        Cách an toàn: xóa text các paragraph trong range, rồi clone format
        paragraph đầu để insert nội dung mới.
        """
        paragraphs = self.doc.paragraphs
        if start_idx >= len(paragraphs) or end_idx > len(paragraphs):
            raise IndexError(f'Range không hợp lệ: {start_idx}-{end_idx}')

        # Lấy paragraph mẫu để clone
        template_p = paragraphs[start_idx]
        # Clone XML element
        template_xml = template_p._p

        # Xóa paragraph từ end_idx-1 lùi về start_idx+1 (giữ paragraph đầu để overwrite)
        for i in range(end_idx - 1, start_idx, -1):
            p = paragraphs[i]
            p._element.getparent().remove(p._element)

        # Sửa paragraph đầu thành new_paragraphs[0]
        if new_paragraphs:
            self._set_paragraph_text(template_p, new_paragraphs[0].get('text', ''))
            # Apply bold/italic nếu có
            if new_paragraphs[0].get('bold'):
                for run in template_p.runs:
                    run.bold = True
            if new_paragraphs[0].get('italic'):
                for run in template_p.runs:
                    run.italic = True

        # Insert các paragraph còn lại sau template_p, clone XML
        prev_element = template_p._element
        for new_p in new_paragraphs[1:]:
            new_xml = deepcopy(template_xml)
            prev_element.addnext(new_xml)
            from docx.text.paragraph import Paragraph
            new_paragraph = Paragraph(new_xml, template_p._parent)
            self._set_paragraph_text(new_paragraph, new_p.get('text', ''))
            if new_p.get('bold'):
                for run in new_paragraph.runs:
                    run.bold = True
                    run.italic = False
            elif new_p.get('italic'):
                for run in new_paragraph.runs:
                    run.bold = False
                    run.italic = True
            else:
                for run in new_paragraph.runs:
                    run.bold = False
                    run.italic = False
            prev_element = new_xml

    # ============================================================
    # DEBUG: liệt kê paragraphs và bảng
    # ============================================================

    def inspect(self):
        """In cấu trúc file để hiểu vị trí các phần."""
        print('=== PARAGRAPHS ===')
        for i, p in enumerate(self.doc.paragraphs):
            text = p.text[:80]
            print(f'  P{i}: [{text}]')
        print('=== TABLES ===')
        for ti, t in enumerate(self.doc.tables):
            print(f'  Table {ti}: {len(t.rows)} hàng × {len(t.columns)} cột')
            for ri, row in enumerate(t.rows):
                for ci, cell in enumerate(row.cells):
                    for pi, p in enumerate(cell.paragraphs):
                        text = p.text[:60]
                        print(f'    [{ri},{ci}] P{pi}: {text}')

    # ============================================================
    # LƯU FILE
    # ============================================================

    def save(self, output_path):
        import os
        os.makedirs(os.path.dirname(output_path), exist_ok=True) if os.path.dirname(output_path) else None
        self.doc.save(output_path)
        print(f'✓ Đã lưu: {output_path}')


if __name__ == '__main__':
    # Demo: đọc cấu trúc file mẫu
    import sys
    if len(sys.argv) > 1:
        TemplateDoc(sys.argv[1]).inspect()
    else:
        print('Cách dùng: python3 fill_template.py templates/01-cong-van.docx')
