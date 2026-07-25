#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_bpb.py — Sinh file .docx BÀI PHÁT BIỂU của Giám đốc Sở Công Thương Lào Cai
đúng định dạng đã duyệt: Times New Roman 15pt, lề T1.5/B1.25/L3.0/R1.25 (cm), tiêu đề
in đậm căn giữa, lời kính thưa in nghiêng thụt lề, thân bài căn đều thụt dòng đầu 1.27cm.

CÁCH DÙNG
---------
    python build_bpb.py noidung.txt "2026.02.15. Bai phat bieu KTCK.docx"

File nội dung dùng markup nhẹ (mỗi dòng bắt đầu bằng 1 thẻ). Trong P/H/LEAD/BREAK/CLOSE
có thể dùng **in đậm** và *in nghiêng* nội dòng.

  [TITLE]      dòng tiêu đề — căn giữa, in đậm, IN HOA giữ nguyên như gõ.
  [GREET-HEAD] dòng "Kính thưa:" đầu — căn đều, in đậm, thụt dòng đầu.
  [GREET]      dòng đối tượng kính thưa — in nghiêng, thụt lề trái 1.27cm.
  [BREAK]      dấu ngắt "Kính thưa Hội nghị!" — in đậm + nghiêng.
  [H]          đề mục nhóm nhiệm vụ (1. / I - / * Nhóm 1:) — in đậm, thụt dòng đầu.
  [LEAD]       đoạn có cụm dẫn đầu in đậm — tự dùng **...** để đánh dấu phần đậm.
  [P]          đoạn văn thường — căn đều, thụt dòng đầu.
  [I]          ý phụ (1)(2)(3) hoặc trích dẫn — in nghiêng, thụt dòng đầu.
  [CLOSE]      dòng cảm ơn/kết — như [P].
  [BLANK]      chèn 1 dòng trống.

Dòng không có thẻ được coi như [P]. Dòng bắt đầu bằng '##' là chú thích, bỏ qua.
"""
import sys, re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

FONT = "Times New Roman"
SIZE = 15          # pt — cỡ chữ bài phát biểu (to hơn VBHC 14pt để đọc trên bục)
INDENT = Cm(1.27)  # thụt dòng đầu / thụt lề lời kính thưa


def _set_run(run, bold=None, italic=None):
    run.font.name = FONT
    run.font.size = Pt(SIZE)
    # đảm bảo font cho cả tiếng Việt (đông á / cs)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = rpr.makeelement(qn('w:rFonts'), {})
        rpr.insert(0, rfonts)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(a), FONT)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _add_inline(p, text, base_bold=False, base_italic=False):
    """Phân tích **đậm** và *nghiêng* nội dòng, thêm các run tương ứng."""
    # tách theo **...** trước, rồi *...*
    token = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
    for tok in token:
        if not tok:
            continue
        b, i = base_bold, base_italic
        s = tok
        if tok.startswith('**') and tok.endswith('**'):
            b = True
            s = tok[2:-2]
        elif tok.startswith('*') and tok.endswith('*'):
            i = True
            s = tok[1:-1]
        r = p.add_run(s)
        _set_run(r, bold=b, italic=i)


def build(lines):
    doc = Document()
    # docDefaults + Normal
    normal = doc.styles['Normal']
    normal.font.name = FONT
    normal.font.size = Pt(SIZE)
    normal.paragraph_format.line_spacing = 1.0
    normal.paragraph_format.space_before = Pt(6)
    normal.paragraph_format.space_after = Pt(6)

    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    sec.top_margin = Cm(1.5)
    sec.bottom_margin = Cm(1.25)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(1.25)

    def newp():
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing = 1.0
        pf.space_before = Pt(6)
        pf.space_after = Pt(6)
        return p

    for raw in lines:
        line = raw.rstrip('\n')
        if line.strip().startswith('##'):
            continue
        m = re.match(r'^\s*\[([A-Z-]+)\]\s?(.*)$', line)
        tag = m.group(1) if m else 'P'
        body = m.group(2) if m else line

        if tag == 'BLANK' or (not body.strip() and tag == 'P'):
            newp()
            continue

        p = newp()
        pf = p.paragraph_format

        if tag == 'TITLE':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_inline(p, body, base_bold=True)
        elif tag == 'GREET-HEAD':
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf.first_line_indent = INDENT
            _add_inline(p, body, base_bold=True)
        elif tag == 'GREET':
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf.left_indent = INDENT
            _add_inline(p, body, base_italic=True)
        elif tag == 'BREAK':
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _add_inline(p, body, base_bold=True, base_italic=True)
        elif tag == 'H':
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf.first_line_indent = INDENT
            _add_inline(p, body, base_bold=True)
        elif tag == 'I':
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf.first_line_indent = INDENT
            _add_inline(p, body, base_italic=True)
        else:  # P, LEAD, CLOSE
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf.first_line_indent = INDENT
            _add_inline(p, body)

    return doc


def main():
    if len(sys.argv) < 3:
        print(__doc__)
        sys.exit(1)
    src, out = sys.argv[1], sys.argv[2]
    with open(src, encoding='utf-8') as f:
        lines = f.readlines()
    doc = build(lines)
    doc.save(out)
    print("Đã lưu:", out)


if __name__ == '__main__':
    main()
