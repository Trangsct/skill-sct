#!/usr/bin/env python3
"""
Kiểm tra văn bản hành chính trước khi trình ký Lãnh đạo Sở/UBND.

Dò 4 nhóm rủi ro:
  A. Bịa số văn bản: ký hiệu văn bản đáng nghi (số trống "Số: /...",
     hoặc số quá lớn/quá nhỏ so với khoảng hợp lý).
  B. Từ suy đoán: "dự kiến", "gần như", "có thể", "khả năng cao"...
  C. Văn bản pháp luật đã hết hiệu lực (danh mục đen có sẵn).
  D. Thiếu phần "cần xác minh" với các thông tin đáng nghi.

Cách dùng:
  python check_document.py <file.docx>
  python check_document.py <file.docx> --strict   # fail nếu có cảnh báo nhóm A/C

Output: in cảnh báo theo từng nhóm, kèm dòng/đoạn cụ thể.
"""

import re
import sys
from pathlib import Path

# ============================================================
# Cấu hình
# ============================================================

# Nhóm B — Từ suy đoán cấm dùng trong văn bản hành chính trình ký
FORBIDDEN_SPECULATIVE = [
    r"\bdự kiến\b",
    r"\bgần như\b",
    r"\bcó vẻ\b",
    r"\bcó lẽ\b",
    r"\bkhả năng cao\b",
    r"\bnhiều khả năng\b",
    r"\bphỏng đoán\b",
    r"\bước chừng\b",
    r"\btheo tôi\b",
    r"\btôi nghĩ\b",
    r"\btôi cho rằng\b",
    r"\bcó thể (?:là|sẽ|không)\b",  # "có thể là", "có thể sẽ" — suy đoán
]

# Cho phép "có thể" trong các ngữ cảnh hợp lệ (tránh false positive)
SPECULATIVE_ALLOWED_CONTEXTS = [
    r"có thể được",  # "có thể được áp dụng" — diễn đạt chấp nhận
    r"có thể tham khảo",
    r"có thể xem xét",
]

# Nhóm C — Danh mục VBQPPL đã hết hiệu lực hoặc đã được thay thế
EXPIRED_LAWS = {
    # Luật
    r"Luật Đầu tư\s+(?:số\s+)?61/2020/QH14": "Đã thay bằng Luật Đầu tư 143/2025/QH15 (HL 01/3/2026)",
    r"Luật Hóa chất\s+(?:năm\s+)?2007": "Đã thay bằng Luật Hóa chất 69/2025/QH15 (HL 01/01/2026)",
    r"Luật Hoá chất\s+(?:năm\s+)?2007": "Đã thay bằng Luật Hóa chất 69/2025/QH15 (HL 01/01/2026)",
    # Nghị định CCN
    r"Nghị định\s+(?:số\s+)?68/2017/NĐ-CP": "Đã thay bằng NĐ 32/2024/NĐ-CP",
    r"Nghị định\s+(?:số\s+)?66/2020/NĐ-CP": "Đã thay bằng NĐ 32/2024/NĐ-CP",
    r"NĐ\s+68/2017": "Đã thay bằng NĐ 32/2024/NĐ-CP",
    r"NĐ\s+66/2020": "Đã thay bằng NĐ 32/2024/NĐ-CP",
    # Thông tư CCN
    r"Thông tư\s+(?:số\s+)?28/2020/TT-BCT": "Đã bãi bỏ bởi TT 14/2024/TT-BCT",
    r"TT\s+28/2020/TT-BCT": "Đã bãi bỏ bởi TT 14/2024/TT-BCT",
    # Quyết định bí mật nhà nước
    r"Quyết định\s+(?:số\s+)?1369/(?:2020/)?QĐ-TTg": "Đã thay bằng QĐ 476/QĐ-TTg ngày 25/3/2026",
    r"QĐ\s+1369/2020/QĐ-TTg": "Đã thay bằng QĐ 476/QĐ-TTg ngày 25/3/2026",
    # Văn bản môi trường cũ
    r"Nghị định\s+(?:số\s+)?40/2019/NĐ-CP": "Đã hết hiệu lực",
    r"Nghị định\s+(?:số\s+)?15/2021/NĐ-CP": "Đã hết hiệu lực, cần kiểm tra văn bản thay thế",
    # Quy chuẩn cũ
    r"QCVN\s+40:2011/BTNMT": "Đã thay bằng QCVN 40:2025/BTNMT (TT 06/2025/TT-BTNMT)",
    # ===== Bổ sung 05/7/2026 (rà soát các phiên làm việc cuối 6 - đầu 7/2026) =====
    # Xây dựng (khung mới hiệu lực 01/7/2026 và 01/01/2026)
    r"Luật Xây dựng\s+(?:số\s+)?50/2014/QH13": "Đã thay bằng Luật Xây dựng 135/2025/QH15",
    r"Nghị định\s+(?:số\s+)?06/2021/NĐ-CP": "Đã thay bằng NĐ 207/2026/NĐ-CP (quản lý chất lượng, thi công, bảo trì)",
    r"Nghị định\s+(?:số\s+)?175/2024/NĐ-CP": "Đã thay bằng NĐ 217/2026/NĐ-CP (quản lý hoạt động xây dựng)",
    r"Thông tư\s+(?:số\s+)?06/2021/TT-BXD": "Đã thay bằng TT 34/2026/TT-BXD (phân cấp công trình)",
    # VLNCN
    r"Nghị định\s+(?:số\s+)?71/2018/NĐ-CP": "Đã thay bằng NĐ 181/2024/NĐ-CP (VLNCN, tiền chất thuốc nổ)",
    # HHNH — phân cấp thẩm quyền
    r"Điều\s+3\s+Thông tư\s+(?:số\s+)?15/2026/TT-BCT": "Đã bị bãi bỏ bởi Điều 26 TT 26/2026/TT-BCT — dẫn Điều 25 TT 26/2026",
    # Hóa chất (khung mới hiệu lực 01/01/2026)
    r"Nghị định\s+(?:số\s+)?113/2017/NĐ-CP": "Đã thay bằng NĐ 24-26/2026/NĐ-CP (khung Luật Hóa chất 69/2025)",
    r"Nghị định\s+(?:số\s+)?82/2022/NĐ-CP": "Đã thay bằng NĐ 24-26/2026/NĐ-CP (khung Luật Hóa chất 69/2025)",
    r"Thông tư\s+(?:số\s+)?32/2017/TT-BCT": "Đã thay bằng TT 01+02/2026/TT-BCT",
    # Thanh tra/kiểm tra
    r"thanh tra chuyên ngành": "Từ Luật Thanh tra 84/2025 các sở KHÔNG còn thanh tra — dùng 'kiểm tra chuyên ngành' theo NĐ 217/2025",
}

# Nhóm A — Pattern số văn bản đáng nghi
SUSPICIOUS_NUMBER_PATTERNS = [
    # Số văn bản rất lớn (có thể là bịa)
    (r"\bsố\s+(\d{5,})\s*\/", "Số văn bản có 5+ chữ số — đáng nghi, cần xác minh"),
]

# Các trạng thái "để trống" HỢP LỆ theo quy ước đã chốt (SKILL.md — mục "Rà soát/review"
# và Nhóm G): số văn bản do văn thư cấp khi phát hành; dòng ngày điền sẵn tháng/năm,
# để trống ngày cho văn thư điền khi ký. CHỈ ghi nhận thông tin, KHÔNG tính cảnh báo
# chặn trình ký. (Sửa 05/7/2026 — trước đây xếp nhầm vào Nhóm A gây false positive.)
INFO_BLANK_PATTERNS = [
    (r"Số\s*:\s*\/[A-ZĐ]", "Số văn bản để trống (đúng quy trình — văn thư cấp khi phát hành)"),
    (r"ngày\s+\/\s*\/\d{4}", "Ngày văn bản để trống (đúng quy trình — điền khi ký)"),
    (r"ngày\s+tháng\s+\d+\s+năm\s+\d{4}", "Ngày trong tháng để trống, tháng/năm đã điền sẵn (đúng chuẩn Nhóm G)"),
]


# ============================================================
# Đọc văn bản
# ============================================================

def read_docx(path: Path) -> str:
    """Đọc nội dung text từ file .docx."""
    try:
        from docx import Document
        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs]
        # Đọc cả nội dung trong table
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)
        return "\n".join(paragraphs)
    except ImportError:
        # Fallback: dùng pandoc nếu có
        import subprocess
        try:
            out = subprocess.check_output(
                ["pandoc", str(path), "-t", "plain"],
                stderr=subprocess.DEVNULL, timeout=30,
            )
            return out.decode("utf-8", errors="replace")
        except (subprocess.CalledProcessError, FileNotFoundError):
            return ""


def read_text(path: Path) -> str:
    """Đọc file text/md."""
    return path.read_text(encoding="utf-8", errors="replace")


def read_input(path: Path) -> str:
    suf = path.suffix.lower()
    if suf == ".docx":
        return read_docx(path)
    if suf in (".txt", ".md"):
        return read_text(path)
    raise ValueError(f"Không hỗ trợ định dạng {suf}. Hỗ trợ: .docx, .txt, .md")


def find_header_linebreaks(path: Path) -> list[tuple[str, str]]:
    """
    Dò lỗi NGẮT DÒNG CỨNG (<w:br/>) lọt vào header/tiêu đề — lỗi tái diễn
    (vd V/v bị gãy 'Kế / hoạch phát / triển...'). Trả về danh sách
    (vị trí, đoạn text) các ô header có <w:br/>.

    Chỉ áp dụng cho .docx. Quét bảng đầu tiên (header) và mọi paragraph
    nằm ngoài bảng ở 1/3 đầu tài liệu (vùng tiêu đề: tên loại VB, trích yếu).
    """
    if path.suffix.lower() != ".docx":
        return []
    try:
        from docx import Document
    except ImportError:
        return []
    hits: list[tuple[str, str]] = []
    doc = Document(str(path))
    # Header table = bảng đầu tiên (UBND/SỞ | Quốc hiệu)
    if doc.tables:
        t0 = doc.tables[0]
        for ri, row in enumerate(t0.rows):
            for ci, cell in enumerate(row.cells):
                for p in cell.paragraphs:
                    if "<w:br/>" in p._p.xml or "<w:br " in p._p.xml:
                        hits.append((f"Header ô[{ri},{ci}]", p.text.strip()[:70]))
    # Vùng tiêu đề ngoài bảng (tên loại VB, trích yếu): 12 paragraph đầu
    for p in doc.paragraphs[:12]:
        if "<w:br/>" in p._p.xml or "<w:br " in p._p.xml:
            hits.append(("Tiêu đề (ngoài bảng)", p.text.strip()[:70]))
    return hits


# ============================================================
# Phát hiện rủi ro
# ============================================================

def find_speculative(text: str) -> list[tuple[int, str, str]]:
    """Tìm các từ/cụm suy đoán. Trả về [(line_no, pattern_matched, snippet)]."""
    warnings = []
    lines = text.splitlines()
    for i, line in enumerate(lines, 1):
        line_lower = line.lower()
        # Kiểm tra các ngữ cảnh được phép
        allowed_spans = []
        for allow_pat in SPECULATIVE_ALLOWED_CONTEXTS:
            for m in re.finditer(allow_pat, line_lower):
                allowed_spans.append((m.start(), m.end()))
        # Tìm pattern cấm
        for pat in FORBIDDEN_SPECULATIVE:
            for m in re.finditer(pat, line_lower):
                # Bỏ qua nếu nằm trong ngữ cảnh được phép
                if any(s <= m.start() < e for s, e in allowed_spans):
                    continue
                snippet = line.strip()
                if len(snippet) > 160:
                    start = max(0, m.start() - 50)
                    end = min(len(snippet), m.end() + 50)
                    snippet = "..." + snippet[start:end] + "..."
                warnings.append((i, pat, snippet))
    return warnings


def find_expired_laws(text: str) -> list[tuple[int, str, str, str]]:
    """Tìm văn bản pháp luật đã hết hiệu lực. Trả về [(line_no, vb_cu, ly_do, snippet)]."""
    warnings = []
    lines = text.splitlines()
    for i, line in enumerate(lines, 1):
        for pat, reason in EXPIRED_LAWS.items():
            m = re.search(pat, line, re.IGNORECASE | re.UNICODE)
            if m:
                snippet = line.strip()
                if len(snippet) > 160:
                    start = max(0, m.start() - 50)
                    end = min(len(snippet), m.end() + 50)
                    snippet = "..." + snippet[start:end] + "..."
                warnings.append((i, m.group(0), reason, snippet))
    return warnings


def find_suspicious_numbers(text: str) -> list[tuple[int, str, str]]:
    """Tìm số văn bản đáng nghi. Trả về [(line_no, ly_do, snippet)]."""
    warnings = []
    lines = text.splitlines()
    for i, line in enumerate(lines, 1):
        for pat, reason in SUSPICIOUS_NUMBER_PATTERNS:
            m = re.search(pat, line, re.IGNORECASE)
            if m:
                snippet = line.strip()
                if len(snippet) > 160:
                    start = max(0, m.start() - 50)
                    end = min(len(snippet), m.end() + 50)
                    snippet = "..." + snippet[start:end] + "..."
                warnings.append((i, reason, snippet))
    return warnings


def find_info_blanks(text: str) -> list[tuple[int, str]]:
    """Ghi nhận các trạng thái 'để trống' HỢP LỆ (chỉ thông tin, không cảnh báo)."""
    infos = []
    lines = text.splitlines()
    for i, line in enumerate(lines, 1):
        for pat, reason in INFO_BLANK_PATTERNS:
            if re.search(pat, line, re.IGNORECASE):
                infos.append((i, reason))
    return infos


def find_unverified_claims(text: str) -> list[tuple[int, str]]:
    """
    Tìm các đoạn có dấu hiệu chưa xác minh nhưng không có ghi chú xác minh.
    Hiện đơn giản: tìm số văn bản đầy đủ mà không có cụm "đề nghị xác minh"
    gần đó trong cùng đoạn.
    """
    warnings = []
    # Đếm tổng số văn bản pháp luật xuất hiện
    law_count = len(re.findall(
        r"(?:Quyết định|Nghị định|Thông tư|Luật|Nghị quyết|Công văn|Kế hoạch)\s+"
        r"(?:số\s+)?\d+",
        text, re.IGNORECASE,
    ))
    verify_count = len(re.findall(
        r"(?:đề nghị|cần|yêu cầu)\s+(?:Bạn\s+)?xác minh", text, re.IGNORECASE,
    ))
    if law_count >= 5 and verify_count == 0:
        warnings.append((
            0,
            f"Văn bản trích dẫn {law_count} văn bản pháp luật nhưng KHÔNG có "
            f"ghi chú 'đề nghị xác minh' nào. Đảm bảo mọi số/ngày đều đã được "
            f"tra cứu, hoặc bổ sung phụ lục xác minh cho các số chưa chắc chắn."
        ))
    return warnings


# ============================================================
# Main
# ============================================================

def main():
    if len(sys.argv) < 2:
        print(__doc__, file=sys.stderr)
        sys.exit(1)
    path = Path(sys.argv[1])
    strict = "--strict" in sys.argv
    if not path.exists():
        print(f"Không tìm thấy file: {path}", file=sys.stderr)
        sys.exit(1)

    text = read_input(path)
    if not text.strip():
        print("Không đọc được nội dung file. Cài python-docx: pip install python-docx",
              file=sys.stderr)
        sys.exit(1)

    line_count = len(text.splitlines())
    print(f"=== KIỂM TRA VĂN BẢN: {path.name} ===")
    print(f"Tổng số dòng: {line_count}\n")

    has_critical = False

    # [INFO] — Trạng thái để trống hợp lệ (không tính cảnh báo)
    infos = find_info_blanks(text)
    if infos:
        print(f"[INFO] Trạng thái để trống hợp lệ ({len(infos)} mục — không tính cảnh báo):")
        for ln, reason in infos[:10]:
            print(f"   Dòng {ln}: {reason}")
        print()

    # Nhóm A — Số văn bản đáng nghi
    sus = find_suspicious_numbers(text)
    if sus:
        print(f"[A] SỐ VĂN BẢN ĐÁNG NGHI ({len(sus)} cảnh báo):")
        for ln, reason, snippet in sus[:20]:
            print(f"   Dòng {ln}: {reason}")
            print(f"      → {snippet}")
        has_critical = True
        print()

    # Nhóm B — Từ suy đoán
    spec = find_speculative(text)
    if spec:
        print(f"[B] TỪ SUY ĐOÁN CẤM DÙNG ({len(spec)} cảnh báo):")
        for ln, pat, snippet in spec[:20]:
            pat_clean = pat.replace(r"\b", "").replace(r"\s+", " ")
            print(f"   Dòng {ln}: pattern \"{pat_clean}\"")
            print(f"      → {snippet}")
        print()

    # Nhóm C — Văn bản hết hiệu lực
    exp = find_expired_laws(text)
    if exp:
        print(f"[C] VĂN BẢN PHÁP LUẬT ĐÃ HẾT HIỆU LỰC ({len(exp)} cảnh báo):")
        for ln, vb, reason, snippet in exp[:20]:
            print(f"   Dòng {ln}: '{vb}' — {reason}")
            print(f"      → {snippet}")
        has_critical = True
        print()

    # Nhóm D — Thiếu ghi chú xác minh
    unv = find_unverified_claims(text)
    if unv:
        print(f"[D] CẢNH BÁO XÁC MINH:")
        for ln, msg in unv:
            print(f"   {msg}")
        print()

    # Nhóm E — Ngắt dòng cứng (<w:br/>) lọt vào header/tiêu đề
    brk = find_header_linebreaks(path)
    if brk:
        print(f"[E] NGẮT DÒNG CỨNG TRONG HEADER/TIÊU ĐỀ ({len(brk)} cảnh báo):")
        for loc, snippet in brk:
            print(f"   {loc}: có <w:br/> — V/v, tên cơ quan hoặc tiêu đề bị ngắt dòng cứng")
            print(f"      → {snippet}")
        print("   Khắc phục: truyền chuỗi LIỀN (không có \\n) cho trường header; để Word tự wrap.")
        has_critical = True
        print()

    # Tổng kết
    total = len(sus) + len(spec) + len(exp) + len(unv) + len(brk)
    if total == 0:
        print("✓ Không phát hiện rủi ro theo các pattern đã định.")
        print("  Lưu ý: script này chỉ dò các pattern bề mặt. Vẫn cần Bạn rà")
        print("  soát thủ công về nội dung pháp lý và phạm vi nhiệm vụ.")
    else:
        print(f"=== TỔNG: {total} cảnh báo ===")
        if has_critical:
            print("⚠ Có cảnh báo NHÓM A, C hoặc E — KHÔNG nên trình ký trước khi xử lý.")

    if strict and has_critical:
        sys.exit(2)


if __name__ == "__main__":
    main()
